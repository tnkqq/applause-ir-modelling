#!/usr/bin/env python3
"""Generate diploma-ready report for experiment series 3 and update docs/Диплом.docx."""

from __future__ import annotations

import json
import math
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import cv2
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

if __package__ is None or __package__ == "":
    sys.path.append(str(Path(__file__).resolve().parents[2]))

from experiments.series3.common import (  # noqa: E402
    ADC_BITS,
    ADC_MAX,
    apply_filter,
    detect_global_threshold,
    generate_adc_frame,
    rng_from_seed,
    save_heatmap,
    scene_with_anomaly,
)


ROOT = Path(__file__).resolve().parents[2]
RESULTS = ROOT / "results"
DOCS = ROOT / "docs"
DOC_IMAGES = DOCS / "images"
REPORTS = RESULTS / "reports"
REPORTS.mkdir(parents=True, exist_ok=True)
DOC_IMAGES.mkdir(parents=True, exist_ok=True)

REPORT_MD = DOCS / "report.md"
REPORT_DOCX = DOCS / "report.docx"
REPORTS_MD = REPORTS / "series3_experiments_report.md"
REPORTS_DOCX = REPORTS / "series3_experiments_report.docx"
SUMMARY_CSV = DOCS / "experiment_summary.csv"
FRAME_STATS_CSV = DOCS / "series3_representative_frame_stats.csv"
DIPLOMA = DOCS / "Диплом.docx"
DIPLOMA_BACKUP = REPORTS / "Диплом_before_series3_update.docx"
DIPLOMA_COPY = REPORTS / "Диплом_with_series3.docx"
REPRESENTATIVE_FRAME_STATS: list[dict[str, Any]] = []


@dataclass(frozen=True)
class ExperimentInfo:
    number: int
    folder: Path
    short: str
    goal: str
    modeled: str
    parameters: str
    interpretation: str
    main_graph: Path


def fmt(value: Any, digits: int = 3) -> str:
    try:
        num = float(value)
    except (TypeError, ValueError):
        return str(value)
    if math.isnan(num):
        return "-"
    if num != 0 and (abs(num) < 1e-3 or abs(num) >= 1e5):
        return f"{num:.{digits}e}"
    return f"{num:.{digits}f}"


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def normalize_u8(arr: np.ndarray) -> np.ndarray:
    data = np.asarray(arr, dtype=float)
    mn, mx = float(np.min(data)), float(np.max(data))
    if math.isclose(mn, mx):
        return np.zeros(data.shape, dtype=np.uint8)
    return np.round(np.clip((data - mn) / (mx - mn), 0, 1) * 255).astype(np.uint8)


def record_frame_stats(number: int, frame: np.ndarray, mask: np.ndarray | None = None) -> None:
    arr = np.asarray(frame, dtype=float)
    row: dict[str, Any] = {
        "run": f"3.{number}",
        "frame_mean_adc": float(np.mean(arr)),
        "frame_std_adc": float(np.std(arr)),
        "frame_min_adc": float(np.min(arr)),
        "frame_max_adc": float(np.max(arr)),
        "frame_dynamic_range_adc": float(np.max(arr) - np.min(arr)),
        "mean_signal_adc": np.nan,
        "mean_background_adc": float(np.mean(arr)),
        "std_background_adc": float(np.std(arr)),
        "snr_like": np.nan,
    }
    if mask is not None and np.any(mask) and np.any(~mask):
        mask_bool = np.asarray(mask, dtype=bool)
        signal = arr[mask_bool]
        background = arr[~mask_bool]
        bg_std = float(np.std(background))
        row["mean_signal_adc"] = float(np.mean(signal))
        row["mean_background_adc"] = float(np.mean(background))
        row["std_background_adc"] = bg_std
        row["snr_like"] = float((np.mean(signal) - np.mean(background)) / bg_std) if bg_std > 0 else np.nan
    REPRESENTATIVE_FRAME_STATS.append(row)


def save_histogram(array: np.ndarray, filename: str, title: str, xlabel: str = "Код ADC") -> str:
    path = DOC_IMAGES / filename
    plt.figure(figsize=(7.2, 4.2))
    plt.hist(np.asarray(array, dtype=float).ravel(), bins=48, color="#3b6ea8", edgecolor="white")
    plt.title(title)
    plt.xlabel(xlabel)
    plt.ylabel("Число пикселей")
    plt.grid(True, axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig(path, dpi=170)
    plt.close()
    return f"images/{filename}"


def save_report_heatmap(array: np.ndarray, filename: str, title: str, label: str = "Код ADC") -> str:
    path = DOC_IMAGES / filename
    save_heatmap(path, array, title, label)
    return f"images/{filename}"


def copy_doc_image(src: Path, filename: str) -> str:
    dst = DOC_IMAGES / filename
    shutil.copy2(src, dst)
    return f"images/{filename}"


def representative_images() -> dict[int, dict[str, str]]:
    """Generate report heatmaps/histograms and copy key experiment graphs."""
    REPRESENTATIVE_FRAME_STATS.clear()
    out: dict[int, dict[str, str]] = {}

    rng = rng_from_seed(3101)
    temp_map = np.full((96, 128), 320.0)
    fpn = rng.normal(0.0, 0.8, size=temp_map.shape)
    frame, _ = generate_adc_frame(temp_map, rng, gaussian_sigma=2.5, fpn_std=0.8, fpn_pattern=fpn)
    record_frame_stats(1, frame)
    out[1] = {
        "heat": save_report_heatmap(frame, "series3_exp01_heatmap_320K.png", "Эксперимент 1: однородный кадр, 320 K"),
        "hist": save_histogram(frame, "series3_exp01_hist_320K.png", "Эксперимент 1: гистограмма, 320 K"),
        "graph": copy_doc_image(RESULTS / "experiment_01_signal_validation" / "signal_vs_temperature.png", "series3_exp01_signal_vs_temperature.png"),
        "extra": copy_doc_image(RESULTS / "experiment_01_signal_validation" / "netd_estimation.png", "series3_exp01_netd_estimation.png"),
    }

    rng = rng_from_seed(3202)
    scene, mask = scene_with_anomaly(96, 128, background_k=300, delta_t=2.0, shape="circle", size=14)
    frame, _ = generate_adc_frame(scene, rng, gaussian_sigma=2.2, fpn_std=0.7)
    record_frame_stats(2, frame, mask)
    out[2] = {
        "heat": save_report_heatmap(frame, "series3_exp02_heatmap_dT2.png", "Эксперимент 2: кадр с аномалией, dT=2 K"),
        "hist": save_histogram(frame, "series3_exp02_hist_dT2.png", "Эксперимент 2: гистограмма, dT=2 K"),
        "mask": save_report_heatmap(mask.astype(float), "series3_exp02_mask_dT2.png", "Эксперимент 2: истинная маска", "Маска"),
        "graph": copy_doc_image(RESULTS / "experiment_02_min_detectable_contrast" / "detection_probability_vs_delta_t.png", "series3_exp02_detection_probability.png"),
    }

    rng = rng_from_seed(3303)
    scene, mask = scene_with_anomaly(96, 128, background_k=300, delta_t=4.0, shape="circle", size=16, weak_bg=True)
    frame, _ = generate_adc_frame(scene, rng, gaussian_sigma=9.0, fpn_std=6.0, quant_bits=5, defect_rate=0.012)
    record_frame_stats(3, frame, mask)
    out[3] = {
        "heat": save_report_heatmap(frame, "series3_exp03_heatmap_combined_noise.png", "Эксперимент 3: комбинированный шум"),
        "hist": save_histogram(frame, "series3_exp03_hist_combined_noise.png", "Эксперимент 3: гистограмма, комбинированный шум"),
        "graph": copy_doc_image(RESULTS / "experiment_03_noise_influence" / "iou_vs_noise_level.png", "series3_exp03_iou_vs_noise_level.png"),
        "extra": copy_doc_image(RESULTS / "experiment_03_noise_influence" / "snr_vs_noise_level.png", "series3_exp03_snr_vs_noise_level.png"),
    }

    rng = rng_from_seed(3404)
    scene, mask = scene_with_anomaly(96, 128, background_k=300, delta_t=3.5, shape="circle", size=14, weak_bg=True)
    raw, _ = generate_adc_frame(scene, rng, gaussian_sigma=4.0, fpn_std=2.0, quant_bits=9, defect_rate=0.002)
    filtered = apply_filter(raw, "median", {"ksize": 3})
    record_frame_stats(4, filtered, mask)
    out[4] = {
        "heat": save_report_heatmap(filtered, "series3_exp04_heatmap_median_k3.png", "Эксперимент 4: медианный фильтр k=3"),
        "hist": save_histogram(filtered, "series3_exp04_hist_median_k3.png", "Эксперимент 4: гистограмма после медианного фильтра k=3"),
        "graph": copy_doc_image(RESULTS / "experiment_04_filtering" / "iou_by_filter.png", "series3_exp04_iou_by_filter.png"),
        "extra": copy_doc_image(RESULTS / "experiment_04_filtering" / "filter_examples.png", "series3_exp04_filter_examples.png"),
    }

    rng = rng_from_seed(3505)
    scene, mask = scene_with_anomaly(96, 128, background_k=300, delta_t=6.0, shape="rectangle", size=8, fill_factor=1.0)
    frame, _ = generate_adc_frame(scene, rng, gaussian_sigma=2.5, fpn_std=0.8)
    record_frame_stats(5, frame, mask)
    out[5] = {
        "heat": save_report_heatmap(frame, "series3_exp05_heatmap_size8_fill1.png", "Эксперимент 5: размер=8, fill=1"),
        "hist": save_histogram(frame, "series3_exp05_hist_size8_fill1.png", "Эксперимент 5: гистограмма, размер=8, fill=1"),
        "mask": save_report_heatmap(mask.astype(float), "series3_exp05_mask_size8_fill1.png", "Эксперимент 5: маска, размер=8", "Маска"),
        "graph": copy_doc_image(RESULTS / "experiment_05_spatial_resolution" / "tpr_heatmap.png", "series3_exp05_tpr_heatmap.png"),
    }

    rng = rng_from_seed(3606)
    scene, mask = scene_with_anomaly(96, 128, background_k=300, delta_t=5.0, shape="circle", size=14, weak_bg=True)
    frame, _ = generate_adc_frame(scene, rng, gaussian_sigma=2.8, fpn_std=1.0)
    record_frame_stats(6, frame, mask)
    out[6] = {
        "heat": save_report_heatmap(frame, "series3_exp06_heatmap_dynamic_frame.png", "Эксперимент 6: репрезентативный кадр последовательности"),
        "hist": save_histogram(frame, "series3_exp06_hist_dynamic_frame.png", "Эксперимент 6: гистограмма репрезентативного кадра"),
        "mask": save_report_heatmap(mask.astype(float), "series3_exp06_mask_dynamic_frame.png", "Эксперимент 6: маска аномалии", "Маска"),
        "graph": copy_doc_image(RESULTS / "experiment_06_temporal_dynamics" / "detection_delay_vs_alpha.png", "series3_exp06_detection_delay_vs_alpha.png"),
        "extra": copy_doc_image(RESULTS / "experiment_06_temporal_dynamics" / "snr_vs_window_size.png", "series3_exp06_snr_vs_window.png"),
    }

    rng = rng_from_seed(3707)
    scene, mask = scene_with_anomaly(96, 128, background_k=300, delta_t=4.0, shape="rectangle", size=14, center=(46, 78), weak_bg=True)
    frame, _ = generate_adc_frame(scene, rng, gaussian_sigma=3.0, fpn_std=1.2, defect_rate=0.002)
    record_frame_stats(7, frame, mask)
    out[7] = {
        "heat": save_report_heatmap(frame, "series3_exp07_heatmap_detector_frame.png", "Эксперимент 7: кадр для сравнения детекторов"),
        "hist": save_histogram(frame, "series3_exp07_hist_detector_frame.png", "Эксперимент 7: гистограмма кадра"),
        "mask": save_report_heatmap(mask.astype(float), "series3_exp07_mask_detector_frame.png", "Эксперимент 7: истинная маска", "Маска"),
        "graph": copy_doc_image(RESULTS / "experiment_07_detector_comparison" / "f1_by_algorithm.png", "series3_exp07_f1_by_algorithm.png"),
        "extra": copy_doc_image(RESULTS / "experiment_07_detector_comparison" / "iou_by_algorithm.png", "series3_exp07_iou_by_algorithm.png"),
    }

    pipeline_src = RESULTS / "reports" / "block_schemes" / "01_obshchaya_shema_modelirovaniya.png"
    if pipeline_src.exists():
        out[0] = {"pipeline": copy_doc_image(pipeline_src, "series3_pipeline.png")}
    return out


def load_experiments() -> list[ExperimentInfo]:
    return [
        ExperimentInfo(
            1,
            RESULTS / "experiment_01_signal_validation",
            "Валидация ИК-сигнала, SNR и NETD",
            "Проверить монотонность цифрового сигнала от температуры сцены и оценить шумовую температурную чувствительность.",
            "Серия однородных кадров без аномалий при температурах 280-360 K с несколькими реализациями шума.",
            "Температуры 280-360 K, шаг 10 K; 16 кадров на уровень; seed 3101; Gaussian noise 2.5 ADC, FPN 0.8 ADC.",
            "Рост температуры приводит к монотонному росту ADC-кода, а NETD задается отношением шума к локальному наклону радиометрической характеристики.",
            RESULTS / "experiment_01_signal_validation" / "signal_vs_temperature.png",
        ),
        ExperimentInfo(
            2,
            RESULTS / "experiment_02_min_detectable_contrast",
            "Минимальный обнаруживаемый температурный контраст",
            "Определить минимальное dT, при котором простая пороговая сегментация устойчиво выделяет аномалию.",
            "Однородный фон 300 K и локальные аномалии формы circle, rectangle и gaussian hot spot.",
            "dT от 0.1 до 10 K; 12 кадров на уровень; критерий успеха IoU > 0.3; seed 3202.",
            "При dT ниже 1 K аномалия практически не выделяется; начиная с 2 K вероятность успешного обнаружения достигает 1.0.",
            RESULTS / "experiment_02_min_detectable_contrast" / "detection_probability_vs_delta_t.png",
        ),
        ExperimentInfo(
            3,
            RESULTS / "experiment_03_noise_influence",
            "Влияние типа и уровня шума",
            "Показать, какие шумовые компоненты сильнее всего ухудшают обнаружение температурной аномалии.",
            "Фиксированная аномалия dT=4 K на слабо неоднородном фоне с Gaussian noise, FPN, квантованием, дефектами и combined noise.",
            "Уровни шума 0, 1, 2, 4, 6, 8; 10 кадров на вариант; seed 3303.",
            "Комбинированный шум оказался наиболее опасным, так как одновременно снижает SNR и нарушает пространственную связность области аномалии.",
            RESULTS / "experiment_03_noise_influence" / "iou_vs_noise_level.png",
        ),
        ExperimentInfo(
            4,
            RESULTS / "experiment_04_filtering",
            "Фильтрация перед обнаружением",
            "Оценить, улучшают ли простые фильтры качество обнаружения на зашумленных ИК-кадрах.",
            "Одинаковый набор зашумленных кадров обработан фильтрами none, Gaussian, median, bilateral и NLM.",
            "24 кадра; dT=3.5 K; seed 3404; единый detector mean/median + k sigma для всех вариантов.",
            "Median k=3 дал лучший баланс между подавлением шума и сохранением локального контраста аномалии.",
            RESULTS / "experiment_04_filtering" / "iou_by_filter.png",
        ),
        ExperimentInfo(
            5,
            RESULTS / "experiment_05_spatial_resolution",
            "Пространственное разрешение и размер аномалии",
            "Проверить влияние размера аномалии и коэффициента заполнения пикселя на обнаружимость.",
            "Прямоугольные аномалии от 1x1 до 32x32 пикселей при разных fill_factor.",
            "Размеры 1, 2, 4, 8, 16, 32 px; fill_factor 0.1, 0.25, 0.5, 0.75, 1.0; dT=6 K; seed 3505.",
            "Малый fill_factor ослабляет полезный сигнал внутри пикселя и может сделать даже геометрически присутствующую аномалию статистически неустойчивой.",
            RESULTS / "experiment_05_spatial_resolution" / "tpr_heatmap.png",
        ),
        ExperimentInfo(
            6,
            RESULTS / "experiment_06_temporal_dynamics",
            "Временная инерционность и усреднение кадров",
            "Показать компромисс между ростом SNR за счет усреднения и задержкой обнаружения динамических аномалий.",
            "Последовательности со статической, появляющейся и движущейся аномалией; модель инерционности первого порядка.",
            "80 кадров; alpha 0.0, 0.3, 0.6, 0.9; окна усреднения 1, 3, 5, 10; seed 3606.",
            "Усреднение полезно для статических объектов, но высокий alpha увеличивает задержку и снижает амплитуду быстро меняющихся событий.",
            RESULTS / "experiment_06_temporal_dynamics" / "detection_delay_vs_alpha.png",
        ),
        ExperimentInfo(
            7,
            RESULTS / "experiment_07_detector_comparison",
            "Сравнение алгоритмов обнаружения",
            "Сравнить интерпретируемые алгоритмы обнаружения и выбрать базовый вариант для практической части диплома.",
            "Общий набор кадров с разными dT, размерами, шумами, положениями и отдельными кадрами без аномалии.",
            "Алгоритмы: global threshold, adaptive local threshold, Otsu, DoG, morphology; seed 3707.",
            "Global threshold k=2.5 оказался лучшим базовым вариантом по F1 и IoU на сформированном наборе, сохранив минимальную вычислительную сложность.",
            RESULTS / "experiment_07_detector_comparison" / "f1_by_algorithm.png",
        ),
    ]


def experiment_result_tables() -> dict[int, pd.DataFrame]:
    tables: dict[int, pd.DataFrame] = {}
    e1 = pd.read_csv(RESULTS / "experiment_01_signal_validation" / "metrics.csv")
    tables[1] = e1[["temperature_K", "mean_signal_adc", "noise_std_adc", "dynamic_range_adc", "snr_for_temp_step", "netd_K"]]
    e2 = pd.read_csv(RESULTS / "experiment_02_min_detectable_contrast" / "metrics_by_delta_t.csv")
    tables[2] = e2[["delta_t_K", "tpr", "fpr", "precision", "iou", "detection_probability", "snr_like"]]
    e3 = pd.read_csv(RESULTS / "experiment_03_noise_influence" / "noise_type_comparison.csv")
    tables[3] = e3[["noise_type", "snr_like_mean", "tpr_mean", "fpr_mean", "iou_mean"]]
    e4 = pd.read_csv(RESULTS / "experiment_04_filtering" / "filter_comparison.csv").sort_values("iou", ascending=False)
    tables[4] = e4[["label", "snr_after", "snr_improvement", "tpr", "fpr", "precision", "iou", "contrast_change"]].head(9)
    e5 = pd.read_csv(RESULTS / "experiment_05_spatial_resolution" / "min_detectable_size.csv")
    tables[5] = e5[["fill_factor", "min_detectable_size_px"]]
    e6 = pd.read_csv(RESULTS / "experiment_06_temporal_dynamics" / "metrics.csv")
    tables[6] = e6.groupby("alpha", as_index=False).agg(
        detection_delay_frames=("detection_delay_frames", "mean"),
        tpr=("tpr", "mean"),
        fpr=("fpr", "mean"),
        snr_like=("snr_like", "mean"),
        peak_amplitude_ratio=("peak_amplitude_ratio", "mean"),
    )
    e7 = pd.read_csv(RESULTS / "experiment_07_detector_comparison" / "algorithm_summary.csv")
    tables[7] = e7[["algorithm", "best_parameter", "precision", "recall", "fpr", "f1", "iou", "runtime_s"]]
    return tables


def md_table(df: pd.DataFrame, max_rows: int = 12) -> str:
    show = df.head(max_rows).copy()
    for col in show.columns:
        if pd.api.types.is_numeric_dtype(show[col]):
            show[col] = show[col].map(lambda x: fmt(x))
    headers = "|" + "|".join(show.columns) + "|"
    sep = "|" + "|".join(["---"] * len(show.columns)) + "|"
    rows = ["|" + "|".join(str(v) for v in row) + "|" for row in show.to_numpy()]
    return "\n".join([headers, sep, *rows])


def build_summary_csv(experiments: list[ExperimentInfo], tables: dict[int, pd.DataFrame]) -> pd.DataFrame:
    rows = []
    for exp in experiments:
        summary = read_json(exp.folder / "summary.json")
        rows.append(
            {
                "run": f"3.{exp.number}",
                "experiment_folder": exp.folder.name,
                "key_parameters": exp.parameters,
                "key_metrics": summary["main_metrics"],
                "main_result": summary["main_result"],
                "main_plot": str(exp.main_graph),
                "interpretation": exp.interpretation,
            }
        )
    df = pd.DataFrame(rows)
    stats = pd.DataFrame(REPRESENTATIVE_FRAME_STATS)
    if not stats.empty:
        stats.to_csv(FRAME_STATS_CSV, index=False)
        stats.to_csv(REPORTS / "series3_representative_frame_stats.csv", index=False)
        df = df.merge(stats, on="run", how="left")
    df.to_csv(SUMMARY_CSV, index=False)
    df.to_csv(REPORTS / "series3_experiment_summary.csv", index=False)
    return df


def build_markdown(experiments: list[ExperimentInfo], tables: dict[int, pd.DataFrame], images: dict[int, dict[str, str]]) -> str:
    summary_all = pd.read_csv(RESULTS / "summary_all_experiments.csv")
    lines: list[str] = []
    lines.append("# Научно-технический отчет по экспериментам серии 3")
    lines.append("")
    lines.append("## 1. Введение")
    lines.append("")
    lines.append(
        "В настоящем разделе приведены результаты третьей серии вычислительных экспериментов, выполненных на основе программной модели инфракрасного датчика. "
        "Назначение моделирования состоит в том, чтобы получить воспроизводимые синтетические ИК-кадры с известной температурной структурой сцены и проверить, как физические и аппаратные факторы влияют на обнаружение температурных аномалий. "
        "Синтетические данные в этом случае полезны потому, что для каждого кадра известны фон, температурный контраст, истинная маска аномалии, параметры шума и условия обработки. Это позволяет рассчитывать метрики не только визуально, но и количественно."
    )
    lines.append("")
    lines.append(
        "Практическая область применения связана с мониторингом температурных аномалий, где требуется отличать локально нагретые области от фона при наличии шумов матрицы, дефектных пикселей, ограниченного пространственного разрешения и временной инерционности сенсора."
    )
    lines.append("")
    lines.append("## 2. Описание математической модели")
    lines.append("")
    lines.append(
        "Модель следует измерительной цепочке инфракрасного сенсора: сначала температура сцены переводится в поток излучения черного тела, затем оптика распределяет мощность по матрице, болометрическая модель задает отклик чувствительных элементов, Readout формирует аналоговый сигнал, ADC переводит его в цифровой код, а NUC уменьшает фиксированную неоднородность. "
        "В серии 3 использована облегченная CPU-реализация этой цепочки: температурно-зависимый сигнал рассчитывался через существующий блок Blackbody, а шумы, дефекты, фильтрация, временная инерционность и детекторы задавались отдельными воспроизводимыми функциями."
    )
    lines.append("")
    lines.append("Спектральная плотность излучения абсолютно черного тела описывается законом Планка:")
    lines.append("")
    lines.append("```math")
    lines.append(r"L_\lambda(T)=\frac{2hc^2}{\lambda^5}\frac{1}{\exp\left(\frac{hc}{\lambda kT}\right)-1}.")
    lines.append("```")
    lines.append("")
    lines.append("Мощность в рабочем диапазоне 8-14 мкм определяется интегрированием по длинам волн и последующим учетом геометрии оптики:")
    lines.append("")
    lines.append("```math")
    lines.append(r"P(x,y)=P_0\left(\frac{f}{\sqrt{f^2+x^2+y^2}}\right)^4.")
    lines.append("```")
    lines.append("")
    lines.append(
        "Для оценки видимости аномалии использовалась SNR-like метрика:"
    )
    lines.append("")
    lines.append("```math")
    lines.append(r"SNR=\frac{\mu_{anom}-\mu_{bg}}{\sigma_{bg}},")
    lines.append("```")
    lines.append("")
    lines.append("где `mu_anom` - среднее значение в области аномалии, `mu_bg` - среднее значение фона, `sigma_bg` - стандартное отклонение фона.")
    lines.append("")
    lines.append("## 3. Архитектура моделирования")
    lines.append("")
    lines.append("```text")
    lines.append("Blackbody -> Optics -> Bolometers -> Readout -> ADC -> NUC -> Detection metrics")
    lines.append("```")
    lines.append("")
    if 0 in images:
        lines.append(f"![Общая схема моделирования]({images[0]['pipeline']})")
        lines.append("")
    lines.append(
        "В экспериментах серии 3 основной акцент сделан на последующих этапах анализа синтетических кадров: расчете SNR/NETD, определении минимального обнаруживаемого контраста, сравнении типов шума, фильтрации, влиянии размера аномалии, временной динамике и сравнении алгоритмов обнаружения."
    )
    lines.append("")
    stats = pd.DataFrame(REPRESENTATIVE_FRAME_STATS)
    if not stats.empty:
        metric_cols = [
            "run",
            "frame_mean_adc",
            "frame_std_adc",
            "frame_min_adc",
            "frame_max_adc",
            "frame_dynamic_range_adc",
            "snr_like",
        ]
        lines.append("Для контроля численных характеристик репрезентативных кадров рассчитаны среднее значение, стандартное отклонение, минимум, максимум, динамический диапазон и SNR-like.")
        lines.append("")
        lines.append(md_table(stats[metric_cols], max_rows=20))
        lines.append("")
    lines.append("## 4. Анализ экспериментов серии 3")
    lines.append("")
    for exp in experiments:
        summary = read_json(exp.folder / "summary.json")
        img = images[exp.number]
        lines.append(f"### 4.{exp.number}. Эксперимент 3.{exp.number}. {exp.short}")
        lines.append("")
        lines.append(f"Цель эксперимента состояла в следующем: {exp.goal}")
        lines.append("")
        lines.append(f"Моделировалось следующее: {exp.modeled}")
        lines.append("")
        lines.append(f"Основные параметры: {exp.parameters}")
        lines.append("")
        lines.append(f"Главный численный результат: {summary['main_result']}")
        lines.append("")
        lines.append(md_table(tables[exp.number]))
        lines.append("")
        if "heat" in img:
            lines.append(f"![Эксперимент {exp.number}: тепловая карта]({img['heat']})")
            lines.append("")
        if "hist" in img:
            lines.append(f"![Эксперимент {exp.number}: гистограмма]({img['hist']})")
            lines.append("")
        if "mask" in img:
            lines.append(f"![Эксперимент {exp.number}: маска]({img['mask']})")
            lines.append("")
        lines.append(f"![Эксперимент {exp.number}: основной график]({img['graph']})")
        lines.append("")
        if "extra" in img:
            lines.append(f"![Эксперимент {exp.number}: дополнительный график]({img['extra']})")
            lines.append("")
        lines.append(f"Интерпретация результата: {exp.interpretation}")
        lines.append("")
    lines.append("## 5. Сравнительный анализ")
    lines.append("")
    lines.append(md_table(summary_all[["experiment_number", "title", "main_result"]], max_rows=20))
    lines.append("")
    lines.append(
        "Сравнение экспериментов показывает, что температура определяет полезный радиометрический сигнал, но сама по себе не гарантирует обнаружимость. "
        "При увеличении температурного контраста SNR-like возрастает, что видно по эксперименту 2: вероятность обнаружения становится равной 1.0 при dT=2 K. "
        "Шумы и дефекты действуют противоположно: комбинированный шум дал средний IoU 0.497, то есть оказался наиболее разрушительным среди рассмотренных факторов. "
        "Фильтрация способна улучшить результат, но только при умеренных параметрах: median k=3 дал IoU 0.965, тогда как чрезмерное сглаживание уменьшает локальный контраст. "
        "Пространственный фактор также существенен: малый fill_factor ослабляет полезный сигнал даже при наличии геометрической аномалии. "
        "Временная обработка повышает SNR для статических объектов, но высокая инерционность alpha=0.9 увеличивает задержку обнаружения до 2.56 кадра."
    )
    lines.append("")
    lines.append("## 6. Выводы")
    lines.append("")
    lines.append(
        "По результатам серии 3 можно сделать вывод, что модель пригодна для инженерного анализа обнаружения температурных аномалий. "
        "Она воспроизводит монотонную зависимость цифрового сигнала от температуры, позволяет оценивать NETD, определять минимальный обнаруживаемый температурный контраст и сравнивать алгоритмы обработки. "
        "Главное ограничение состоит в том, что серия 3 использует облегченные CPU-модели шумов и считывания, а не полный медленный расчет `Readout` через `fsolve` для каждого пикселя. Это осознанный компромисс: он сохраняет физический смысл факторов и делает эксперименты достаточно быстрыми для воспроизводимого дипломного исследования. "
        "Для практического обнаружения в рамках данной серии наиболее устойчивым базовым алгоритмом оказался глобальный порог `global_threshold` с параметром `k=2.5`, давший F1=0.872 и IoU=0.964 на сформированном наборе кадров."
    )
    lines.append("")
    return "\n".join(lines)


def add_table_to_doc(doc: Document, df: pd.DataFrame, max_rows: int = 12) -> None:
    show = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(show.columns):
        table.rows[0].cells[idx].text = str(col)
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(show.columns):
            value = row[col]
            if isinstance(value, (float, int, np.floating, np.integer)):
                cells[idx].text = fmt(value)
            else:
                cells[idx].text = str(value)


def add_picture_to_doc(doc: Document, rel_path: str, width: float = 5.9) -> None:
    path = DOCS / rel_path
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))


def build_report_docx(experiments: list[ExperimentInfo], tables: dict[int, pd.DataFrame], images: dict[int, dict[str, str]]) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    doc.add_heading("Научно-технический отчет по экспериментам серии 3", level=0)
    doc.add_heading("1. Введение", level=1)
    doc.add_paragraph(
        "В отчете приведены результаты третьей серии вычислительных экспериментов, выполненных на основе программной модели инфракрасного датчика. "
        "Серия направлена на проверку формирования синтетических ИК-кадров и оценку факторов, влияющих на обнаружение температурных аномалий."
    )
    doc.add_heading("2. Математическая модель", level=1)
    doc.add_paragraph(
        "Модель следует цепочке Blackbody -> Optics -> Bolometers -> Readout -> ADC -> NUC. "
        "Спектральное излучение рассчитывается по закону Планка, мощность на пикселе корректируется оптическим множителем cos^4, а выходной сигнал анализируется через цифровой ADC-код и метрики обнаружения."
    )
    doc.add_paragraph(r"Формулы: L_lambda(T)=2hc^2/lambda^5 /(exp(hc/(lambda kT))-1); SNR=(mu_anom-mu_bg)/sigma_bg.")
    doc.add_heading("3. Архитектура моделирования", level=1)
    doc.add_paragraph("Blackbody -> Optics -> Bolometers -> Readout -> ADC -> NUC -> Detection metrics")
    if 0 in images:
        add_picture_to_doc(doc, images[0]["pipeline"], width=3.8)
    stats = pd.DataFrame(REPRESENTATIVE_FRAME_STATS)
    if not stats.empty:
        doc.add_paragraph(
            "Для репрезентативных кадров каждого эксперимента рассчитаны mean, std, min, max, dynamic range и SNR-like."
        )
        add_table_to_doc(
            doc,
            stats[
                [
                    "run",
                    "frame_mean_adc",
                    "frame_std_adc",
                    "frame_min_adc",
                    "frame_max_adc",
                    "frame_dynamic_range_adc",
                    "snr_like",
                ]
            ],
            max_rows=20,
        )
    doc.add_heading("4. Анализ экспериментов серии 3", level=1)
    for exp in experiments:
        summary = read_json(exp.folder / "summary.json")
        doc.add_heading(f"4.{exp.number}. Эксперимент 3.{exp.number}. {exp.short}", level=2)
        doc.add_paragraph(f"Цель эксперимента: {exp.goal}")
        doc.add_paragraph(f"Моделировалось: {exp.modeled}")
        doc.add_paragraph(f"Параметры: {exp.parameters}")
        doc.add_paragraph(f"Результат: {summary['main_result']}")
        add_table_to_doc(doc, tables[exp.number])
        img = images[exp.number]
        for key in ["heat", "hist", "mask", "graph", "extra"]:
            if key in img:
                add_picture_to_doc(doc, img[key])
        doc.add_paragraph(f"Интерпретация: {exp.interpretation}")
    doc.add_heading("5. Сравнительный анализ", level=1)
    summary_all = pd.read_csv(RESULTS / "summary_all_experiments.csv")
    add_table_to_doc(doc, summary_all[["experiment_number", "title", "main_result"]], max_rows=20)
    doc.add_paragraph(
        "Наиболее сильными факторами оказались температурный контраст, комбинированный шум, пространственный размер аномалии и параметры временной обработки. "
        "NUC и фильтрация повышают пригодность изображения к анализу, но не заменяют контроль дефектов и корректный выбор порога обнаружения."
    )
    doc.add_heading("6. Выводы", level=1)
    doc.add_paragraph(
        "Серия 3 подтверждает пригодность модели для генерации синтетических ИК-данных и проверки алгоритмов обнаружения температурных аномалий. "
        "Лучший базовый детектор в проведенном сравнении - global_threshold с k=2.5; минимальный устойчивый контраст в эксперименте 2 составил 2 K; лучшая оценка NETD в эксперименте 1 составила 0.239 K."
    )
    doc.save(REPORT_DOCX)
    shutil.copy2(REPORT_DOCX, REPORTS_DOCX)


def insert_paragraph_before(anchor, text: str = "", style: str | None = None):
    new_p = anchor.insert_paragraph_before(text)
    if style:
        try:
            new_p.style = style
        except KeyError:
            pass
    return new_p


def move_last_paragraph_before(doc: Document, anchor) -> None:
    elem = doc.paragraphs[-1]._p
    anchor._p.addprevious(elem)


def move_last_table_before(doc: Document, anchor) -> None:
    elem = doc.tables[-1]._tbl
    anchor._p.addprevious(elem)


def add_docx_picture_before(doc: Document, anchor, image_path: Path, width: float = 5.5) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(image_path), width=Inches(width))
    move_last_paragraph_before(doc, anchor)


def add_docx_table_before(doc: Document, anchor, df: pd.DataFrame, max_rows: int = 8) -> None:
    show = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"
    for idx, col in enumerate(show.columns):
        table.rows[0].cells[idx].text = str(col)
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(show.columns):
            value = row[col]
            cells[idx].text = fmt(value) if isinstance(value, (float, int, np.floating, np.integer)) else str(value)
    move_last_table_before(doc, anchor)


def paragraph_index_by_xml(paragraphs, target) -> int:
    for idx, paragraph in enumerate(paragraphs):
        if paragraph._p is target._p:
            return idx
    return -1


def update_diploma_docx(experiments: list[ExperimentInfo], tables: dict[int, pd.DataFrame], images: dict[int, dict[str, str]]) -> None:
    if not DIPLOMA.exists():
        return
    if not DIPLOMA_BACKUP.exists():
        shutil.copy2(DIPLOMA, DIPLOMA_BACKUP)
    doc = Document(DIPLOMA_BACKUP)
    paragraphs = doc.paragraphs

    heading_map = {}
    for i, para in enumerate(paragraphs):
        txt = " ".join(para.text.split())
        for exp in experiments:
            if txt.startswith(f"4.1.{exp.number}"):
                heading_map[exp.number] = para
                para.text = f"4.1.{exp.number} Задача {exp.number}. {exp.short}"

    # Re-read paragraphs after heading text edits.
    paragraphs = doc.paragraphs
    next_anchor = {}
    for exp in experiments:
        current = heading_map.get(exp.number)
        if current is None:
            continue
        cur_idx = paragraph_index_by_xml(paragraphs, current)
        if cur_idx < 0:
            continue
        anchor = None
        for cand in paragraphs[cur_idx + 1 :]:
            style_name = cand.style.name if cand.style else ""
            text = " ".join(cand.text.split())
            if style_name.startswith("Heading") and (text.startswith("4.1.") or text.startswith("Глава 7")):
                anchor = cand
                break
        if anchor is None:
            anchor = paragraphs[-1]
        next_anchor[exp.number] = anchor

    for exp in experiments:
        anchor = next_anchor.get(exp.number)
        if anchor is None:
            continue
        summary = read_json(exp.folder / "summary.json")
        insert_paragraph_before(anchor, f"Цель эксперимента. {exp.goal}", "Body Text")
        insert_paragraph_before(anchor, f"Моделируемая сцена. {exp.modeled}", "Body Text")
        insert_paragraph_before(anchor, f"Параметры моделирования. {exp.parameters}", "Body Text")
        insert_paragraph_before(anchor, f"Полученные результаты. {summary['main_result']}", "Body Text")
        add_docx_table_before(doc, anchor, tables[exp.number], max_rows=6)
        for key in ["heat", "hist", "mask", "graph", "extra"]:
            if key in images[exp.number]:
                add_docx_picture_before(doc, anchor, DOCS / images[exp.number][key], width=5.3)
        insert_paragraph_before(anchor, f"Инженерная интерпретация. {exp.interpretation}", "Body Text")

    doc.save(DIPLOMA)
    shutil.copy2(DIPLOMA, DIPLOMA_COPY)


def main() -> None:
    experiments = load_experiments()
    tables = experiment_result_tables()
    images = representative_images()
    build_summary_csv(experiments, tables)
    md = build_markdown(experiments, tables, images)
    REPORT_MD.write_text(md, encoding="utf-8")
    REPORTS_MD.write_text(md, encoding="utf-8")
    build_report_docx(experiments, tables, images)
    update_diploma_docx(experiments, tables, images)
    print(f"Generated {REPORTS_DOCX}")
    print(f"Generated {REPORT_MD}")
    print(f"Updated {DIPLOMA}")


if __name__ == "__main__":
    main()
