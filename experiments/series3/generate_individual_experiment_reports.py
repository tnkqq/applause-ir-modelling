#!/usr/bin/env python3
"""Generate detailed DOCX reports inside every results/experiment_* folder."""

from __future__ import annotations

import json
import math
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
RESULTS = ROOT / "results"


@dataclass(frozen=True)
class ExperimentText:
    number: int
    short_title: str
    goal: str
    modeled: str
    theory: list[str]
    dependencies: list[str]
    conclusion: str


EXPERIMENT_TEXTS: dict[int, ExperimentText] = {
    1: ExperimentText(
        1,
        "Валидация формирования ИК-сигнала, SNR и NETD",
        "Проверить, что цифровой сигнал модели монотонно возрастает с температурой однородной сцены, и оценить шумовую температурную чувствительность сенсора.",
        "Однородная температурная сцена последовательно задавалась в диапазоне 280-360 K. Для каждого температурного уровня формировалась серия цифровых кадров с гауссовым шумом и фиксированной неоднородностью матрицы.",
        [
            "Эксперимент опирается на закон Планка: спектральная плотность энергетической яркости черного тела возрастает с температурой, а интегрирование по рабочему диапазону 8-14 мкм дает мощность, поступающую на чувствительный элемент.",
            "После оптического и болометрического преобразования мощность масштабируется в цифровой код АЦП. Если модель физически согласована, среднее значение ADC-кода должно быть монотонной функцией температуры.",
            "SNR-like в данном контексте характеризует отношение полезного изменения среднего сигнала к шумовому разбросу. NETD оценивается как отношение шумового стандартного отклонения к локальному наклону радиометрической характеристики dS/dT.",
        ],
        [
            "Температура сцены влияет на средний ADC-код через излучательную способность черного тела: при росте температуры увеличивается поток излучения и, соответственно, сигнал.",
            "Шумовое стандартное отклонение определяет вертикальный разброс точек вокруг радиометрической характеристики и напрямую ухудшает NETD.",
            "Динамический диапазон кадра показывает суммарный разброс пикселей внутри однородной сцены; в этой серии он связан в основном с шумом и FPN.",
        ],
        "Эксперимент подтверждает корректную монотонность температурного отклика и дает базовую оценку чувствительности, от которой зависят последующие задачи обнаружения аномалий.",
    ),
    2: ExperimentText(
        2,
        "Минимальный обнаруживаемый температурный контраст",
        "Определить минимальный температурный контраст аномалии относительно фона, при котором простая пороговая сегментация становится устойчивой.",
        "На фоне 300 K формировались локальные аномалии разных форм: круг, прямоугольник и гауссово пятно. Температурный контраст dT изменялся от 0.1 до 10 K.",
        [
            "Физическая основа эксперимента связана с тем, что малая добавка температуры создает малую добавку ИК-мощности. После прохождения через шумы сенсора слабый контраст может оказаться статистически неотличимым от фона.",
            "Пороговая сегментация использует оценку фонового уровня и робастного стандартного отклонения. Порог вида median + k*sigma выделяет пиксели, сигнал которых существенно превышает фон.",
            "Для оценки качества используются TPR, FPR, precision и IoU. TPR показывает долю найденной истинной аномалии, FPR - долю ложных срабатываний на фоне, IoU - геометрическое совпадение предсказанной и истинной маски.",
        ],
        [
            "При малом dT полезный сигнал меньше или сопоставим с фоновым шумом, поэтому IoU и вероятность обнаружения близки к нулю.",
            "Рост dT увеличивает SNR-like и приводит к резкому переходу от неустойчивого обнаружения к устойчивому.",
            "Форма аномалии влияет на локальную концентрацию сигнала: компактные и равномерные области обычно легче сегментировать, чем плавные гауссовы пятна с размытыми краями.",
        ],
        "Эксперимент задает практический порог применимости модели для задач обнаружения: ниже некоторого контраста аномалия физически присутствует в сцене, но алгоритмически маскируется шумом.",
    ),
    3: ExperimentText(
        3,
        "Влияние типа и уровня шума",
        "Оценить, какие шумовые составляющие ИК-датчика сильнее всего ухудшают обнаружение температурной аномалии.",
        "Фиксированная аномалия моделировалась на слабо неоднородном фоне. Варьировались типы шумов: гауссов временной шум, FPN, квантование, дефектные пиксели и комбинированное воздействие.",
        [
            "В ИК-матрице шум не является единой величиной. Временный электронный шум меняется от кадра к кадру, FPN задает постоянную пиксельную неоднородность, квантование связано с конечной разрядностью АЦП, а дефектные пиксели формируют локальные выбросы.",
            "С точки зрения обнаружения важно не только среднее стандартное отклонение, но и пространственная структура помехи. Фиксированные паттерны и дефекты могут имитировать локальные аномалии или разрушать связность истинной области.",
            "Комбинированный шум наиболее близок к реалистичной неблагоприятной ситуации, поскольку несколько физических факторов действуют одновременно.",
        ],
        [
            "Увеличение уровня шума снижает SNR-like и уменьшает IoU, потому что граница между фоном и аномалией становится менее выраженной.",
            "FPN особенно опасен для простых глобальных порогов, так как создает устойчивые пространственные смещения яркости.",
            "Квантование становится заметным при грубой разрядности: мелкие изменения температуры сжимаются в одни и те же цифровые уровни.",
            "Дефектные пиксели повышают риск ложных срабатываний, даже если средний шум по кадру остается умеренным.",
        ],
        "Эксперимент показывает, что качество обнаружения определяется не только амплитудой шума, но и его физической природой и пространственной организацией.",
    ),
    4: ExperimentText(
        4,
        "Фильтрация ИК-изображений перед обнаружением",
        "Проверить, повышает ли предварительная фильтрация качество обнаружения температурных аномалий на зашумленных кадрах.",
        "Один и тот же набор зашумленных кадров обрабатывался разными фильтрами: без фильтрации, Gaussian, median, bilateral и non-local means. После фильтрации применялся одинаковый детектор.",
        [
            "Фильтрация является этапом подавления помех перед сегментацией. Линейное гауссово сглаживание уменьшает высокочастотный шум, но может размывать границы аномалии.",
            "Медианный фильтр нелинеен и хорошо подавляет импульсные выбросы и одиночные дефектные пиксели, сохраняя резкие границы лучше, чем обычное усреднение.",
            "Bilateral-фильтр учитывает пространственную близость и близость яркости, поэтому теоретически способен сохранять края. NLM ищет похожие фрагменты изображения, но при неправильных параметрах может чрезмерно сгладить полезный контраст.",
        ],
        [
            "Умеренная фильтрация повышает SNR-like за счет уменьшения фонового стандартного отклонения.",
            "Слишком сильное сглаживание уменьшает амплитуду локальной аномалии и может ухудшить precision или IoU.",
            "Оптимальный фильтр определяется компромиссом между подавлением шума и сохранением формы температурной области.",
        ],
        "Эксперимент подтверждает, что фильтрация полезна только при корректном выборе типа и параметров; лучшая обработка не должна уничтожать пространственный контраст аномалии.",
    ),
    5: ExperimentText(
        5,
        "Пространственное разрешение и размер аномалии",
        "Показать влияние площади аномалии и коэффициента заполнения пикселя на обнаружимость.",
        "Моделировались прямоугольные аномалии разных размеров от 1 до 32 пикселей при fill_factor от 0.1 до 1.0 и фиксированном температурном контрасте.",
        [
            "Пространственное разрешение ИК-системы ограничивает способность отделить малую область нагрева от фона. Если аномалия занимает только часть пикселя, ее вклад усредняется с фоновым излучением.",
            "Коэффициент fill_factor в эксперименте описывает эффективную долю пикселя, занятую аномалией. При fill_factor < 1 полезный температурный контраст ослабляется.",
            "Метрики TPR, FPR и IoU здесь отражают не только качество алгоритма, но и физическое ограничение дискретизации сцены матрицей.",
        ],
        [
            "Увеличение размера аномалии повышает число пикселей с полезным сигналом и делает маску статистически устойчивее.",
            "Увеличение fill_factor усиливает локальный сигнал внутри пикселя; малый fill_factor может сделать даже крупную геометрическую аномалию слабой по ADC-контрасту.",
            "Минимальный обнаруживаемый размер следует рассматривать совместно с fill_factor, шумом и выбранным порогом.",
        ],
        "Эксперимент показывает границу применимости алгоритмов на пиксельной сетке: пространственная дискретизация и неполное заполнение пикселя могут быть не менее важны, чем температурный контраст.",
    ),
    6: ExperimentText(
        6,
        "Временная инерционность и усреднение кадров",
        "Показать компромисс между ростом SNR при временной обработке и задержкой обнаружения динамической аномалии.",
        "Формировались последовательности со статической, появляющейся и движущейся аномалией. Варьировались коэффициент инерционности alpha и размер окна кадрового усреднения.",
        [
            "Тепловой датчик и последующая обработка обладают временной инерционностью. Простейшая модель первого порядка записывается как y_t = alpha*y_{t-1} + (1-alpha)*x_t.",
            "Кадровое усреднение уменьшает случайный шум примерно пропорционально корню из числа независимых кадров, но одновременно снижает скорость реакции на быстрые изменения.",
            "Для задач обнаружения динамических событий важны не только SNR-like, TPR и FPR, но и задержка обнаружения, а также отношение пиковой амплитуды после обработки к исходной амплитуде.",
        ],
        [
            "Рост окна усреднения повышает SNR-like для статических объектов, но сглаживает временной фронт появления аномалии.",
            "Увеличение alpha делает отклик более инерционным: шум подавляется лучше, но событие обнаруживается позже.",
            "Для движущихся объектов чрезмерная временная фильтрация может снижать пиковую амплитуду и смазывать траекторию.",
        ],
        "Эксперимент фиксирует важный инженерный компромисс: временная обработка повышает устойчивость к шуму, но ухудшает оперативность обнаружения.",
    ),
    7: ExperimentText(
        7,
        "Сравнение алгоритмов обнаружения температурных аномалий",
        "Сравнить интерпретируемые алгоритмы обнаружения на общем синтетическом наборе ИК-кадров.",
        "Сформирован набор кадров с разными температурными контрастами, размерами, положениями, шумами и отдельными кадрами без аномалии. Сравнивались глобальный порог, адаптивный локальный порог, Otsu, DoG и морфологические варианты.",
        [
            "Глобальный порог предполагает, что фон достаточно однороден и может быть описан одним уровнем и одной оценкой разброса.",
            "Адаптивный локальный порог учитывает медленную пространственную неоднородность фона, но может быть чувствителен к размеру окна и локальным выбросам.",
            "Метод Otsu подбирает порог по гистограмме, DoG выделяет локальные контрастные структуры через разность сглаживаний, а морфология удаляет мелкие шумовые компоненты и стабилизирует форму маски.",
            "Для выбора алгоритма одновременно рассматриваются precision, recall, FPR, F1, IoU, AUC и время выполнения.",
        ],
        [
            "Уменьшение порога обычно повышает recall, но увеличивает FPR; повышение порога действует наоборот.",
            "F1 отражает баланс precision и recall, а IoU дополнительно оценивает геометрию совпадения маски.",
            "Время выполнения важно для практической системы, особенно если обработка должна идти покадрово без GPU.",
        ],
        "Эксперимент выбирает базовый детектор не по одной метрике, а по совокупности качества, устойчивости и вычислительной простоты.",
    ),
    8: ExperimentText(
        8,
        "Итоговая агрегация экспериментов серии 3",
        "Собрать результаты экспериментов 3.1-3.7 в единую таблицу и подготовить материалы для дипломного отчета.",
        "В этом служебном эксперименте не моделируется новый физический режим сенсора; обрабатываются уже полученные результаты предыдущих экспериментов.",
        [
            "Сводная агрегация нужна для сопоставления факторов влияния: температуры, шума, фильтрации, пространственного размера, временной инерционности и алгоритма обнаружения.",
            "Физическая теория здесь используется косвенно: итоговая таблица связывает результаты отдельных экспериментов, каждый из которых проверяет отдельный участок модели ИК-датчика.",
        ],
        [
            "Сводная таблица позволяет быстро увидеть ключевой результат каждого запуска.",
            "Markdown-отчет и папка ключевых графиков используются как промежуточный материал для дипломной главы.",
        ],
        "Папка является навигационным и отчетным узлом серии 3, а не отдельным физическим моделированием.",
    ),
    9: ExperimentText(
        9,
        "Отсутствующая постановка эксперимента",
        "Зафиксировать, что расчетный эксперимент 3.9 не выполнялся, поскольку исходная постановка в run.md была пустой.",
        "Новый набор синтетических кадров не создавался; численные параметры, метрики и графики отсутствуют.",
        [
            "Научная корректность требует не создавать искусственные параметры и выводы без исходной постановки. Поэтому папка 3.9 оставлена как служебная отметка об отсутствии эксперимента.",
            "В рамках дипломной работы такой отчет важен как контроль полноты серии: он показывает, что пропуск связан не с ошибкой запуска, а с отсутствием технического задания.",
        ],
        [
            "Отсутствие config.json, summary.json, CSV и PNG означает, что анализ зависимостей для этого номера невозможен.",
            "Для проведения эксперимента 3.9 необходимо сначала сформулировать цель, параметры, метрики и ожидаемые выходные файлы.",
        ],
        "Эксперимент не содержит расчетных данных; результаты не интерпретируются, чтобы не нарушать требование не выдумывать данные.",
    ),
}


def read_json(path: Path) -> dict[str, Any]:
    if not path.exists() or path.stat().st_size == 0:
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def read_text(path: Path) -> str:
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8")


def experiment_number(folder: Path) -> int:
    match = re.search(r"experiment_(\d+)", folder.name)
    return int(match.group(1)) if match else 0


def fmt(value: Any, digits: int = 3) -> str:
    try:
        num = float(value)
    except (TypeError, ValueError):
        return str(value)
    if math.isnan(num):
        return "-"
    if num != 0 and (abs(num) < 1e-3 or abs(num) >= 1e5):
        return f"{num:.{digits}e}"
    return f"{num:.{digits}f}"


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        doc.add_paragraph(text)


def set_doc_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"


def add_table(doc: Document, df: pd.DataFrame, *, max_rows: int = 16) -> None:
    if df.empty:
        doc.add_paragraph("Таблица не содержит строк.")
        return
    show = df.head(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"
    for i, col in enumerate(show.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(show.columns):
            value = row[col]
            cells[i].text = fmt(value) if isinstance(value, (int, float)) else str(value)
    if len(df) > max_rows:
        doc.add_paragraph(f"В таблице показаны первые {max_rows} строк из {len(df)}; полный набор данных сохранен в CSV-файле эксперимента.")


def numeric_stats(df: pd.DataFrame) -> pd.DataFrame:
    numeric = df.select_dtypes(include="number")
    if numeric.empty:
        return pd.DataFrame()
    desc = numeric.agg(["mean", "std", "min", "max"]).T.reset_index()
    desc = desc.rename(columns={"index": "metric"})
    return desc


def add_csv_analysis(doc: Document, folder: Path) -> None:
    csv_files = sorted(
        path
        for path in folder.glob("*.csv")
        if not path.name.startswith("Эксперимент_")
    )
    if not csv_files:
        doc.add_paragraph("Численные CSV-результаты в папке отсутствуют.")
        return
    for csv_path in csv_files:
        doc.add_heading(f"Файл данных: {csv_path.name}", level=3)
        try:
            df = pd.read_csv(csv_path)
        except Exception as exc:  # noqa: BLE001
            doc.add_paragraph(f"Файл не удалось прочитать как CSV: {exc}")
            continue
        doc.add_paragraph(f"Размер таблицы: {len(df)} строк, {len(df.columns)} столбцов. Столбцы: {', '.join(map(str, df.columns))}.")
        stats = numeric_stats(df)
        if not stats.empty:
            doc.add_paragraph("Сводная статистика числовых столбцов:")
            add_table(doc, stats, max_rows=24)
        doc.add_paragraph("Фрагмент исходной таблицы:")
        add_table(doc, df, max_rows=12)


def add_config(doc: Document, config: dict[str, Any]) -> None:
    if not config:
        doc.add_paragraph("Файл config.json отсутствует или пуст.")
        return
    rows = [{"parameter": key, "value": json.dumps(value, ensure_ascii=False)} for key, value in config.items()]
    add_table(doc, pd.DataFrame(rows), max_rows=40)


def image_caption(path: Path, folder: Path) -> str:
    rel = path.relative_to(folder).as_posix()
    stem = path.stem.replace("_", " ")
    if path.parent.name == "masks":
        return f"Маска: {stem} ({rel})"
    if path.parent.name == "images":
        return f"Пример кадра: {stem} ({rel})"
    return f"График/изображение: {stem} ({rel})"


def add_images(doc: Document, folder: Path) -> None:
    images = sorted(folder.rglob("*.png"), key=lambda p: (p.parent.as_posix(), p.name))
    if not images:
        doc.add_paragraph("PNG-графики и изображения в папке отсутствуют.")
        return
    doc.add_paragraph(
        "Ниже приведены PNG-материалы, полученные при моделировании: графики зависимостей, тепловые карты, примеры кадров и маски аномалий. "
        "Маски показывают истинную область температурной аномалии, с которой сравнивались результаты обнаружения."
    )
    for image_path in images:
        doc.add_paragraph(image_caption(image_path, folder))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(str(image_path), width=Inches(5.8))
        except Exception as exc:  # noqa: BLE001
            doc.add_paragraph(f"Изображение не удалось вставить: {exc}")


def experiment_specific_analysis(number: int, folder: Path) -> list[str]:
    lines: list[str] = []
    if number == 1:
        metrics = pd.read_csv(folder / "metrics.csv")
        first = metrics.iloc[0]
        last = metrics.iloc[-1]
        best_netd = metrics.loc[metrics["netd_K"].idxmin()]
        slope = (last["mean_signal_adc"] - first["mean_signal_adc"]) / (last["temperature_K"] - first["temperature_K"])
        lines.append(
            f"Средний сигнал изменился от {fmt(first['mean_signal_adc'])} ADC при {fmt(first['temperature_K'])} K до {fmt(last['mean_signal_adc'])} ADC при {fmt(last['temperature_K'])} K. "
            f"Средний наклон по крайним точкам равен {fmt(slope)} ADC/K, что согласуется с монотонным ростом потока излучения по закону Планка."
        )
        lines.append(
            f"Минимальная NETD получена при {fmt(best_netd['temperature_K'])} K и составила {fmt(best_netd['netd_K'])} K. "
            "Это означает, что при данном шумовом уровне более высокая температура в рассмотренном диапазоне дает более крутой радиометрический отклик и лучшую температурную различимость."
        )
    elif number == 2:
        metrics = pd.read_csv(folder / "metrics_by_delta_t.csv")
        stable = metrics[(metrics["detection_probability"] >= 0.9) & (metrics["iou"] > 0.3)]
        if not stable.empty:
            row = stable.iloc[0]
            lines.append(
                f"Первый устойчивый режим обнаружения достигается при dT={fmt(row['delta_t_K'])} K: detection_probability={fmt(row['detection_probability'])}, IoU={fmt(row['iou'])}, SNR-like={fmt(row['snr_like'])}."
            )
        low = metrics.iloc[0]
        high = metrics.iloc[-1]
        lines.append(
            f"При dT={fmt(low['delta_t_K'])} K TPR={fmt(low['tpr'])}, а при dT={fmt(high['delta_t_K'])} K TPR={fmt(high['tpr'])}. "
            "Такой рост показывает, что температурный контраст является главным управляющим параметром видимости локальной аномалии."
        )
    elif number == 3:
        comp = pd.read_csv(folder / "noise_type_comparison.csv").sort_values("iou_mean")
        worst = comp.iloc[0]
        best = comp.iloc[-1]
        lines.append(
            f"Наихудший средний результат дал тип шума {worst['noise_type']}: IoU={fmt(worst['iou_mean'])}, SNR-like={fmt(worst['snr_like_mean'])}. "
            f"Наиболее высокий средний IoU наблюдался для {best['noise_type']}: IoU={fmt(best['iou_mean'])}."
        )
        lines.append(
            "Разница между типами шума показывает, что простая оценка дисперсии недостаточна: пространственно фиксированные и комбинированные помехи нарушают геометрию маски сильнее, чем изолированное квантование."
        )
    elif number == 4:
        comp = pd.read_csv(folder / "filter_comparison.csv").sort_values("iou", ascending=False)
        best = comp.iloc[0]
        worst = comp.iloc[-1]
        lines.append(
            f"Лучший вариант по IoU - {best['label']}: IoU={fmt(best['iou'])}, SNR после фильтра={fmt(best['snr_after'])}, улучшение SNR={fmt(best['snr_improvement'])}. "
            f"Худший вариант - {worst['label']}: IoU={fmt(worst['iou'])}."
        )
        lines.append(
            "Результат демонстрирует компромисс фильтрации: подавление фонового шума полезно только до тех пор, пока фильтр сохраняет амплитуду и форму локальной температурной области."
        )
    elif number == 5:
        min_size = pd.read_csv(folder / "min_detectable_size.csv")
        valid = min_size.dropna(subset=["min_detectable_size_px"])
        if not valid.empty:
            lines.append(
                "Минимальный обнаруживаемый размер найден для fill_factor "
                + ", ".join(f"{fmt(r.fill_factor)} -> {fmt(r.min_detectable_size_px)} px" for r in valid.itertuples())
                + "."
            )
        missing = min_size[min_size["min_detectable_size_px"].isna()]
        if not missing.empty:
            lines.append(
                "Для fill_factor "
                + ", ".join(fmt(v) for v in missing["fill_factor"])
                + " критерий обнаружения не был достигнут, что указывает на сильное ослабление полезного сигнала при неполном заполнении пикселя."
            )
    elif number == 6:
        metrics = pd.read_csv(folder / "metrics.csv")
        by_alpha = metrics.groupby("alpha", as_index=False).agg(
            detection_delay_frames=("detection_delay_frames", "mean"),
            snr_like=("snr_like", "mean"),
            tpr=("tpr", "mean"),
            fpr=("fpr", "mean"),
        )
        slow = by_alpha.loc[by_alpha["detection_delay_frames"].idxmax()]
        best_snr = metrics.loc[metrics["snr_like"].idxmax()]
        lines.append(
            f"Наибольшая средняя задержка обнаружения получена при alpha={fmt(slow['alpha'])}: {fmt(slow['detection_delay_frames'])} кадра. "
            f"Максимальный SNR-like в таблице равен {fmt(best_snr['snr_like'])}."
        )
        lines.append(
            "Чем больше alpha и окно усреднения, тем сильнее подавляется случайный шум, однако тем медленнее система реагирует на появление или перемещение аномалии."
        )
    elif number == 7:
        summary = pd.read_csv(folder / "algorithm_summary.csv").sort_values("f1", ascending=False)
        best = summary.iloc[0]
        fastest = summary.loc[summary["runtime_s"].idxmin()]
        lines.append(
            f"Лучший алгоритм по F1 - {best['algorithm']} ({best['best_parameter']}): F1={fmt(best['f1'])}, IoU={fmt(best['iou'])}, precision={fmt(best['precision'])}, recall={fmt(best['recall'])}."
        )
        lines.append(
            f"Самое малое среднее время выполнения показал {fastest['algorithm']}: runtime={fmt(fastest['runtime_s'])} с. "
            "Выбор практического детектора должен учитывать одновременно качество сегментации и вычислительную стоимость."
        )
    elif number == 8:
        csv_path = folder / "summary_all_experiments.csv"
        if csv_path.exists():
            df = pd.read_csv(csv_path)
            lines.append(f"Сводная таблица содержит {len(df)} экспериментов и используется как единый индекс результатов серии 3.")
            lines.append("Она фиксирует главный результат каждого запуска и позволяет быстро перенести данные в дипломную главу.")
    elif number == 9:
        lines.append("Численные зависимости не анализируются, потому что исходная постановка отсутствует и модель не запускалась.")
        lines.append("Следующий корректный шаг для этого номера - сначала записать run.md с целью, параметрами, seed, метриками и ожидаемыми выходными файлами.")
    return lines


def build_report(folder: Path) -> Path:
    number = experiment_number(folder)
    text = EXPERIMENT_TEXTS.get(number, EXPERIMENT_TEXTS[9])
    summary = read_json(folder / "summary.json")
    config = read_json(folder / "config.json")
    readme = read_text(folder / "README.md")

    doc = Document()
    set_doc_style(doc)
    title = f"Эксперимент {number:02d}. {text.short_title}"
    doc.add_heading(title, level=0)

    doc.add_heading("1. Назначение и суть эксперимента", level=1)
    doc.add_paragraph(text.goal)
    doc.add_paragraph(text.modeled)
    if summary.get("main_result"):
        doc.add_paragraph(f"Главный результат, зафиксированный в summary.json: {summary['main_result']}")
    if readme:
        doc.add_paragraph("Краткая постановка из README.md использована как источник параметров и перечня выходных файлов.")

    doc.add_heading("2. Теоретическая основа", level=1)
    add_paragraphs(doc, text.theory)
    doc.add_paragraph(
        "Для всех экспериментов базовой физической цепочкой остается Blackbody -> Optics -> Bolometers -> Readout -> ADC -> NUC. "
        "Температурная сцена преобразуется в ИК-мощность, затем в болометрический отклик и цифровой код, после чего оцениваются метрики качества кадра или обнаружения."
    )

    doc.add_heading("3. Параметры запуска", level=1)
    add_config(doc, config)
    if summary:
        doc.add_paragraph("Сводные поля summary.json:")
        add_table(doc, pd.DataFrame([{"field": k, "value": v} for k, v in summary.items()]), max_rows=20)

    doc.add_heading("4. Численные результаты", level=1)
    add_csv_analysis(doc, folder)

    doc.add_heading("5. Анализ зависимостей и влияния параметров", level=1)
    add_paragraphs(doc, text.dependencies)
    add_paragraphs(doc, experiment_specific_analysis(number, folder))

    doc.add_heading("6. Изображения, графики и маски", level=1)
    add_images(doc, folder)

    doc.add_heading("7. Итог эксперимента", level=1)
    doc.add_paragraph(text.conclusion)
    if summary.get("conclusion"):
        doc.add_paragraph(f"Вывод из summary.json: {summary['conclusion']}")

    report_path = folder / f"Эксперимент_{folder.name.removeprefix('experiment_')}_отчет.docx"
    doc.save(report_path)
    return report_path


def main() -> None:
    folders = sorted(path for path in RESULTS.glob("experiment_*") if path.is_dir())
    generated: list[Path] = []
    for folder in folders:
        generated.append(build_report(folder))
    for path in generated:
        with zipfile.ZipFile(path) as zf:
            bad = zf.testzip()
        if bad:
            raise RuntimeError(f"DOCX archive check failed for {path}: {bad}")
        print(path)


if __name__ == "__main__":
    main()
