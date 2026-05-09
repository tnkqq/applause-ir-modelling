#!/usr/bin/env python3
"""Generate diploma-style report from results/1.x and results/2.x data."""

from __future__ import annotations

import csv
import json
import math
import os
import shutil
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parents[2]
RESULTS = ROOT / "results"
DOCS = ROOT / "docs"
IMAGES = DOCS / "images"
MPLCONFIG = RESULTS / ".mplconfig"
os.environ["MPLCONFIGDIR"] = str(MPLCONFIG)
MPLCONFIG.mkdir(parents=True, exist_ok=True)

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import numpy as np


def read_csv(path: Path) -> list[dict[str, Any]]:
    with path.open(newline="", encoding="utf-8") as file:
        rows = list(csv.DictReader(file))
    return [{key: parse_value(value) for key, value in row.items()} for row in rows]


def parse_value(value: str) -> Any:
    if value is None:
        return None
    text = str(value).strip()
    if text == "":
        return None
    try:
        if any(ch in text.lower() for ch in [".", "e"]):
            return float(text)
        return int(text)
    except ValueError:
        return text


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def write_csv(path: Path, rows: list[dict[str, Any]], fieldnames: list[str]) -> None:
    with path.open("w", newline="", encoding="utf-8") as file:
        writer = csv.DictWriter(file, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def load_array(path: Path) -> np.ndarray:
    if path.suffix == ".npy":
        return np.load(path)
    return np.loadtxt(path, delimiter=",")


def fmt(value: Any, digits: int = 3) -> str:
    if value is None:
        return "—"
    if isinstance(value, str):
        return value
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if math.isnan(number):
        return "—"
    if number != 0 and (abs(number) < 1e-3 or abs(number) >= 1e5):
        return f"{number:.{digits}e}"
    return f"{number:.{digits}f}"


def md_table(rows: list[dict[str, Any]], headers: list[tuple[str, str]]) -> str:
    if not rows:
        return ""
    out = ["|" + "|".join(title for _, title in headers) + "|"]
    out.append("|" + "|".join("---" for _ in headers) + "|")
    for row in rows:
        out.append("|" + "|".join(str(row.get(key, "—")) for key, _ in headers) + "|")
    return "\n".join(out)


def reset_images() -> None:
    IMAGES.mkdir(parents=True, exist_ok=True)
    for path in IMAGES.glob("*.png"):
        path.unlink()


def save_heatmap(array: np.ndarray, filename: str, title: str, label: str) -> str:
    path = IMAGES / filename
    plt.figure(figsize=(7.2, 4.6))
    plt.imshow(array, cmap="inferno")
    plt.colorbar(label=label)
    plt.title(title)
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    return f"images/{filename}"


def save_histogram(array: np.ndarray, filename: str, title: str, xlabel: str) -> str:
    path = IMAGES / filename
    data = np.asarray(array, dtype=float).ravel()
    plt.figure(figsize=(7.2, 4.2))
    plt.hist(data, bins=48, color="#3b6ea8", edgecolor="white")
    plt.title(title)
    plt.xlabel(xlabel)
    plt.ylabel("Число пикселей")
    plt.grid(True, axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    return f"images/{filename}"


def save_line(xs, ys, filename: str, title: str, xlabel: str, ylabel: str, labels=None) -> str:
    path = IMAGES / filename
    plt.figure(figsize=(7.2, 4.2))
    if labels is None:
        plt.plot(xs, ys, marker="o")
    else:
        for x, y, label in zip(xs, ys, labels):
            plt.plot(x, y, marker="o", label=label)
        plt.legend()
    plt.title(title)
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    return f"images/{filename}"


def save_bar(labels, values, filename: str, title: str, xlabel: str, ylabel: str) -> str:
    path = IMAGES / filename
    plt.figure(figsize=(7.2, 4.2))
    plt.bar(labels, values, color="#3b6ea8")
    plt.title(title)
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.grid(True, axis="y", alpha=0.3)
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    return f"images/{filename}"


def save_dual_heatmap(left: np.ndarray, right: np.ndarray, filename: str, title_left: str, title_right: str, label: str) -> str:
    path = IMAGES / filename
    vmin = min(float(np.min(left)), float(np.min(right)))
    vmax = max(float(np.max(left)), float(np.max(right)))
    fig, axes = plt.subplots(1, 2, figsize=(10, 4.2))
    image = axes[0].imshow(left, cmap="inferno", vmin=vmin, vmax=vmax)
    axes[0].set_title(title_left)
    axes[1].imshow(right, cmap="inferno", vmin=vmin, vmax=vmax)
    axes[1].set_title(title_right)
    fig.colorbar(image, ax=axes.ravel().tolist(), label=label, shrink=0.85)
    plt.savefig(path, dpi=180, bbox_inches="tight")
    plt.close()
    return f"images/{filename}"


def save_failure_figure(filename: str) -> str:
    path = IMAGES / filename
    plt.figure(figsize=(8, 3))
    plt.axis("off")
    text = (
        "Запуск 1.2 завершился диагностической ошибкой\\n"
        "Readout.process(): fsolve возвращает массив,\\n"
        "а код присваивает его скалярной ячейке NumPy."
    )
    plt.text(0.02, 0.6, text, fontsize=12, va="center")
    plt.tight_layout()
    plt.savefig(path, dpi=180)
    plt.close()
    return f"images/{filename}"


def summarize_by(rows: list[dict[str, Any]], key: str, value_keys: list[str]) -> list[dict[str, Any]]:
    result = []
    for group in sorted({row[key] for row in rows}):
        subset = [row for row in rows if row[key] == group]
        out = {key: group}
        for value_key in value_keys:
            values = [float(row[value_key]) for row in subset if row.get(value_key) is not None]
            out[f"{value_key}_mean"] = float(np.mean(values)) if values else None
        result.append(out)
    return result


def build_images() -> dict[str, str]:
    reset_images()
    images: dict[str, str] = {}

    rows_10 = read_csv(RESULTS / "1.0" / "blackbody_power.csv")
    images["1.0_line"] = save_line(
        [row["temperature_K"] for row in rows_10],
        [row["power_W"] for row in rows_10],
        "run_1_0_power_vs_temperature.png",
        "Запуск 1.0: мощность на пикселе",
        "Температура, K",
        "P, W",
    )
    images["1.0_hist"] = save_histogram(
        np.array([row["power_W"] for row in rows_10]),
        "run_1_0_power_histogram.png",
        "Запуск 1.0: распределение рассчитанных мощностей",
        "P, W",
    )

    arr_11 = load_array(RESULTS / "1.1" / "arrays" / "p_distribution_400K.csv")
    images["1.1_heat"] = save_heatmap(arr_11, "run_1_1_heatmap_400K.png", "Запуск 1.1: P distribution, 400 K", "W")
    images["1.1_hist"] = save_histogram(arr_11, "run_1_1_hist_400K.png", "Запуск 1.1: histogram P, 400 K", "W")

    images["1.2_failure"] = save_failure_figure("run_1_2_failure.png")

    arr_13 = load_array(RESULTS / "1.3" / "arrays" / "adc_320K.csv")
    images["1.3_heat"] = save_heatmap(arr_13, "run_1_3_adc_320K.png", "Запуск 1.3: ADC, 320 K", "ADC code")
    images["1.3_hist"] = save_histogram(arr_13, "run_1_3_adc_320K_hist.png", "Запуск 1.3: histogram ADC, 320 K", "ADC code")
    rows_13 = read_csv(RESULTS / "1.3" / "adc_metrics.csv")
    images["1.3_line"] = save_line(
        [row["temperature_K"] for row in rows_13],
        [row["mean_adc"] for row in rows_13],
        "run_1_3_mean_adc.png",
        "Запуск 1.3: средний ADC-код",
        "Температура, K",
        "mean ADC",
    )

    raw_13 = load_array(RESULTS / "1.3" / "arrays" / "adc_320K.csv")
    nuc_14 = load_array(RESULTS / "1.4" / "arrays" / "nuc_320K.csv")
    images["1.4_heat"] = save_dual_heatmap(raw_13, nuc_14, "run_1_4_nuc_320K.png", "До NUC", "После NUC", "ADC code")
    images["1.4_hist"] = save_histogram(nuc_14, "run_1_4_nuc_320K_hist.png", "Запуск 1.4: histogram после NUC", "ADC code")

    arr_20 = load_array(RESULTS / "2.0" / "arrays" / "adc_frame_400K.csv")
    rows_20 = read_csv(RESULTS / "2.0" / "adc_metrics.csv")
    images["2.0_heat"] = save_heatmap(arr_20, "run_2_0_adc_400K.png", "Запуск 2.0: ADC, 400 K", "ADC code")
    images["2.0_hist"] = save_histogram(arr_20, "run_2_0_adc_400K_hist.png", "Запуск 2.0: histogram ADC, 400 K", "ADC code")
    images["2.0_line"] = save_line(
        [row["temperature_K"] for row in rows_20],
        [row["mean"] for row in rows_20],
        "run_2_0_mean_adc_vs_temperature.png",
        "Запуск 2.0: mean ADC vs temperature",
        "Температура, K",
        "mean ADC",
    )

    arr_21 = load_array(RESULTS / "2.1" / "arrays" / "p_distribution_300K_wide.csv")
    rows_21 = read_csv(RESULTS / "2.1" / "optics_metrics.csv")
    images["2.1_heat"] = save_heatmap(arr_21, "run_2_1_p_300K_wide.png", "Запуск 2.1: P distribution, wide FOV", "W")
    images["2.1_hist"] = save_histogram(arr_21, "run_2_1_p_300K_wide_hist.png", "Запуск 2.1: histogram P, wide FOV", "W")
    rows_21_300 = [row for row in rows_21 if row["temperature_K"] == 300]
    images["2.1_bar"] = save_bar(
        [row["fov_case"] for row in rows_21_300],
        [row["corner_drop_percent"] for row in rows_21_300],
        "run_2_1_corner_drop_by_fov.png",
        "Запуск 2.1: падение в углу при 300 K",
        "FOV",
        "corner drop, %",
    )

    arr_22 = load_array(RESULTS / "2.2" / "arrays" / "adc_frame_T380_tol_1em04_seed_101.csv")
    rows_22 = read_csv(RESULTS / "2.2" / "fpn_metrics.csv")
    agg_22 = summarize_by(rows_22, "tolerance", ["std", "peak_to_peak"])
    images["2.2_heat"] = save_heatmap(arr_22, "run_2_2_adc_high_tolerance.png", "Запуск 2.2: ADC, tolerance 1e-4", "ADC code")
    images["2.2_hist"] = save_histogram(arr_22, "run_2_2_adc_high_tolerance_hist.png", "Запуск 2.2: histogram ADC", "ADC code")
    images["2.2_line"] = save_line(
        [row["tolerance"] for row in agg_22],
        [row["std_mean"] for row in agg_22],
        "run_2_2_fpn_std_vs_tolerance.png",
        "Запуск 2.2: FPN STD vs tolerance",
        "Tolerance",
        "mean STD, ADC",
    )

    arr_23 = load_array(RESULTS / "2.3" / "arrays" / "adc_frame_defect_rate_2em02_seed_11.csv")
    rows_23 = read_csv(RESULTS / "2.3" / "defect_metrics.csv")
    agg_23 = summarize_by(rows_23, "defect_rate_requested", ["std", "outlier_count_mean_pm_3std"])
    images["2.3_heat"] = save_heatmap(arr_23, "run_2_3_adc_defects_2pct.png", "Запуск 2.3: ADC с 2% дефектов", "ADC code")
    images["2.3_hist"] = save_histogram(arr_23, "run_2_3_adc_defects_2pct_hist.png", "Запуск 2.3: histogram ADC", "ADC code")
    images["2.3_line"] = save_line(
        [row["defect_rate_requested"] for row in agg_23],
        [row["std_mean"] for row in agg_23],
        "run_2_3_std_vs_defects.png",
        "Запуск 2.3: STD vs defect rate",
        "Доля дефектов",
        "mean STD, ADC",
    )

    unc_24 = load_array(RESULTS / "2.4" / "arrays" / "adc_uncorrected_340K.csv")
    cor_24 = load_array(RESULTS / "2.4" / "arrays" / "adc_corrected_340K_full.csv")
    rows_24 = read_csv(RESULTS / "2.4" / "nuc_metrics.csv")
    images["2.4_heat"] = save_dual_heatmap(unc_24, cor_24, "run_2_4_nuc_340K.png", "До NUC", "После NUC", "ADC code")
    images["2.4_hist"] = save_histogram(cor_24, "run_2_4_corrected_340K_hist.png", "Запуск 2.4: histogram после NUC", "ADC code")
    rows_24_340 = [row for row in rows_24 if row["temperature_K"] == 340]
    images["2.4_bar"] = save_bar(
        [row["fractional_bits"] for row in rows_24_340],
        [row["std_after"] for row in rows_24_340],
        "run_2_4_std_after_by_precision.png",
        "Запуск 2.4: STD после NUC, 340 K",
        "Точность коэффициентов",
        "STD after",
    )

    scene_25 = load_array(RESULTS / "2.5" / "arrays" / "scene_map_tol_low_size_large_Ta_400K.csv")
    unc_25 = load_array(RESULTS / "2.5" / "arrays" / "adc_uncorrected_tol_low_size_large_Ta_400K.csv")
    cor_25 = load_array(RESULTS / "2.5" / "arrays" / "adc_corrected_tol_low_size_large_Ta_400K.csv")
    rows_25 = read_csv(RESULTS / "2.5" / "anomaly_metrics.csv")
    images["2.5_scene"] = save_heatmap(scene_25, "run_2_5_scene_large_400K.png", "Запуск 2.5: карта сцены", "K")
    images["2.5_heat"] = save_dual_heatmap(unc_25, cor_25, "run_2_5_anomaly_400K.png", "До NUC", "После NUC", "ADC code")
    images["2.5_hist"] = save_histogram(cor_25, "run_2_5_corrected_400K_hist.png", "Запуск 2.5: histogram corrected", "ADC code")
    contrast_series = []
    labels = []
    for state in ["uncorrected", "corrected"]:
        subset = [
            row
            for row in rows_25
            if row["tolerance_case"] == "low" and row["anomaly_size"] == "large" and row["state"] == state
        ]
        contrast_series.append(([row["anomaly_temperature_K"] for row in subset], [row["contrast"] for row in subset]))
        labels.append(state)
    images["2.5_line"] = save_line(
        [item[0] for item in contrast_series],
        [item[1] for item in contrast_series],
        "run_2_5_contrast_vs_temperature.png",
        "Запуск 2.5: contrast vs anomaly temperature",
        "Температура аномалии, K",
        "Контраст, ADC",
        labels=labels,
    )

    return images


def collect_data() -> dict[str, Any]:
    return {
        "1.0": {"power": read_csv(RESULTS / "1.0" / "blackbody_power.csv"), "metrics": read_json(RESULTS / "1.0" / "metrics.json")},
        "1.1": {"metrics": read_csv(RESULTS / "1.1" / "optics_metrics.csv")},
        "1.2": {"status": read_json(RESULTS / "1.2" / "status.json")},
        "1.3": {"metrics": read_csv(RESULTS / "1.3" / "adc_metrics.csv")},
        "1.4": {"metrics": read_csv(RESULTS / "1.4" / "nuc_metrics.csv")},
        "2.0": {"adc": read_csv(RESULTS / "2.0" / "adc_metrics.csv"), "power": read_csv(RESULTS / "2.0" / "blackbody_power.csv")},
        "2.1": {"metrics": read_csv(RESULTS / "2.1" / "optics_metrics.csv")},
        "2.2": {"metrics": read_csv(RESULTS / "2.2" / "fpn_metrics.csv")},
        "2.3": {"metrics": read_csv(RESULTS / "2.3" / "defect_metrics.csv")},
        "2.4": {"metrics": read_csv(RESULTS / "2.4" / "nuc_metrics.csv")},
        "2.5": {"metrics": read_csv(RESULTS / "2.5" / "anomaly_metrics.csv")},
    }


def build_experiment_summary(data: dict[str, Any]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    p10 = data["1.0"]["metrics"]
    rows.append(
        {
            "run": "1.0",
            "key_parameters": "Blackbody, 280..420 K, step 20 K",
            "key_metrics": f"P_min={fmt(p10['power_min_W'])} W; P_max={fmt(p10['power_max_W'])} W; ratio={fmt(p10['power_ratio_max_to_min'])}",
            "interpretation": "Проверена монотонная температурная зависимость входной ИК-мощности.",
        }
    )
    r11 = data["1.1"]["metrics"]
    rows.append(
        {
            "run": "1.1",
            "key_parameters": "Blackbody + Optics, 64x48, 300/350/400 K",
            "key_metrics": f"corner/center at 400 K={fmt(r11[-1]['corner_to_center_ratio'])}",
            "interpretation": "Оптика формирует спад мощности к краям кадра.",
        }
    )
    rows.append(
        {
            "run": "1.2",
            "key_parameters": "Small full chain, 8x6",
            "key_metrics": "failed at Readout.process",
            "interpretation": "Выявлена совместимость: fsolve возвращает массив, старый код присваивает его скаляру.",
        }
    )
    r13 = data["1.3"]["metrics"]
    rows.append(
        {
            "run": "1.3",
            "key_parameters": "Blackbody -> Optics -> Bolometers -> Readout -> ADC, 8x6",
            "key_metrics": f"mean ADC 300 K={fmt(r13[0]['mean_adc'])}; mean ADC 320 K={fmt(r13[1]['mean_adc'])}",
            "interpretation": "Полный малый конвейер дал отрицательные ADC-коды, что выявило необходимость физического clip/offset.",
        }
    )
    r14 = data["1.4"]["metrics"]
    rows.append(
        {
            "run": "1.4",
            "key_parameters": "NUC on run 1.3 frames",
            "key_metrics": f"coef_a nonfinite={r14[0]['nonfinite_count']}; coef_b nonfinite={r14[1]['nonfinite_count']}",
            "interpretation": "NUC выполнилась, но наследовала отрицательный диапазон входного ADC.",
        }
    )
    r20 = data["2.0"]["adc"]
    rows.append(
        {
            "run": "2.0",
            "key_parameters": "160x120, 280..400 K, calibrated ADC",
            "key_metrics": f"mean ADC {fmt(r20[0]['mean'])}->{fmt(r20[-1]['mean'])}; sat={fmt(r20[-1]['saturated_fraction'])}",
            "interpretation": "Получена базовая рабочая радиометрическая зависимость без насыщения.",
        }
    )
    r21 = data["2.1"]["metrics"]
    worst21 = max(r21, key=lambda row: row["corner_drop_percent"])
    rows.append(
        {
            "run": "2.1",
            "key_parameters": "160x120, 300/340/380/420 K, narrow/medium/wide FOV",
            "key_metrics": f"max corner drop={fmt(worst21['corner_drop_percent'])}% ({worst21['fov_case']})",
            "interpretation": "Оптика является источником плавной пространственной неравномерности.",
        }
    )
    r22agg = summarize_by(data["2.2"]["metrics"], "tolerance", ["std", "peak_to_peak"])
    worst22 = max(r22agg, key=lambda row: row["std_mean"])
    rows.append(
        {
            "run": "2.2",
            "key_parameters": "64x48, flat field, tolerances 1e-6..1e-4, 3 seeds",
            "key_metrics": f"max mean STD={fmt(worst22['std_mean'])} at tol={worst22['tolerance']}",
            "interpretation": "Разброс параметров болометров проявляется как fixed pattern noise.",
        }
    )
    r23agg = summarize_by(data["2.3"]["metrics"], "defect_rate_requested", ["std", "outlier_count_mean_pm_3std"])
    rows.append(
        {
            "run": "2.3",
            "key_parameters": "96x72, hot/cold defects 0..2%, 5 seeds",
            "key_metrics": f"STD {fmt(r23agg[0]['std_mean'])}->{fmt(r23agg[-1]['std_mean'])}; outliers at 2%={fmt(r23agg[-1]['outlier_count_mean_pm_3std_mean'])}",
            "interpretation": "Дефекты увеличивают дисперсию и число выбросов.",
        }
    )
    r24 = data["2.4"]["metrics"]
    best24 = min(r24, key=lambda row: row["std_after"])
    rows.append(
        {
            "run": "2.4",
            "key_parameters": "96x72, two-point NUC, 3/4/5/6/full fractional precision",
            "key_metrics": f"best std_after={fmt(best24['std_after'])} at {best24['temperature_K']} K/{best24['fractional_bits']}",
            "interpretation": "NUC снижает FPN; разрядность коэффициентов влияет на остаточную неоднородность.",
        }
    )
    r25 = data["2.5"]["metrics"]
    best25 = max(r25, key=lambda row: row["snr_like"])
    rows.append(
        {
            "run": "2.5",
            "key_parameters": "96x72, background 300 K, anomalies 320..400 K, sizes small/medium/large",
            "key_metrics": f"max SNR-like={fmt(best25['snr_like'])} at {best25['anomaly_temperature_K']} K/{best25['anomaly_size']}/{best25['state']}",
            "interpretation": "Сформирована прикладная оценка обнаружимости температурной аномалии.",
        }
    )
    return rows


def report_text(data: dict[str, Any], images: dict[str, str], summary_rows: list[dict[str, str]]) -> str:
    r10 = data["1.0"]["metrics"]
    r11 = data["1.1"]["metrics"]
    r13 = data["1.3"]["metrics"]
    r14 = data["1.4"]["metrics"]
    r20 = data["2.0"]["adc"]
    r21 = data["2.1"]["metrics"]
    r22agg = summarize_by(data["2.2"]["metrics"], "tolerance", ["std", "peak_to_peak"])
    r23agg = summarize_by(data["2.3"]["metrics"], "defect_rate_requested", ["std", "outlier_count_mean_pm_3std"])
    r24 = data["2.4"]["metrics"]
    r25 = data["2.5"]["metrics"]
    best24 = min(r24, key=lambda row: row["std_after"])
    best25 = max(r25, key=lambda row: row["snr_like"])
    worst21 = max(r21, key=lambda row: row["corner_drop_percent"])
    worst22 = max(r22agg, key=lambda row: row["std_mean"])

    table_20 = md_table(
        [
            {
                "T": row["temperature_K"],
                "mean": fmt(row["mean"]),
                "std": fmt(row["std"]),
                "minmax": f"{fmt(row['min'])}/{fmt(row['max'])}",
                "dr": fmt(row["dynamic_range"]),
                "sat": fmt(row["saturated_fraction"]),
            }
            for row in r20
        ],
        [("T", "T, K"), ("mean", "mean"), ("std", "std"), ("minmax", "min/max"), ("dr", "dynamic range"), ("sat", "sat. fraction")],
    )
    table_21 = md_table(
        [
            {
                "fov": row["fov_case"],
                "edge": fmt(row["edge_drop_percent"]),
                "corner": fmt(row["corner_drop_percent"]),
                "radial": fmt(row["radial_nonuniformity_index"]),
            }
            for row in r21
            if row["temperature_K"] == 300
        ],
        [("fov", "FOV"), ("edge", "edge drop, %"), ("corner", "corner drop, %"), ("radial", "radial index")],
    )
    table_22 = md_table(
        [
            {"tol": row["tolerance"], "std": fmt(row["std_mean"]), "p2p": fmt(row["peak_to_peak_mean"])}
            for row in r22agg
        ],
        [("tol", "tolerance"), ("std", "mean std"), ("p2p", "mean peak-to-peak")],
    )
    table_23 = md_table(
        [
            {"rate": row["defect_rate_requested"], "std": fmt(row["std_mean"]), "outliers": fmt(row["outlier_count_mean_pm_3std_mean"])}
            for row in r23agg
        ],
        [("rate", "defect rate"), ("std", "mean std"), ("outliers", "mean outliers")],
    )
    table_24 = md_table(
        [
            {
                "bits": row["fractional_bits"],
                "before": fmt(row["std_before"]),
                "after": fmt(row["std_after"]),
                "res": fmt(row["residual_nonuniformity"]),
            }
            for row in r24
            if row["temperature_K"] == 340
        ],
        [("bits", "precision"), ("before", "std before"), ("after", "std after"), ("res", "std_after/std_before")],
    )
    table_25 = md_table(
        [
            {
                "T": row["anomaly_temperature_K"],
                "size": row["anomaly_size"],
                "state": row["state"],
                "contrast": fmt(row["contrast"]),
                "snr": fmt(row["snr_like"]),
            }
            for row in r25
            if row["tolerance_case"] == "low" and row["anomaly_size"] == "large"
        ],
        [("T", "T anomaly, K"), ("size", "size"), ("state", "state"), ("contrast", "contrast"), ("snr", "SNR-like")],
    )
    summary_table = md_table(
        summary_rows,
        [("run", "Запуск"), ("key_parameters", "Ключевые параметры"), ("key_metrics", "Ключевые метрики"), ("interpretation", "Интерпретация")],
    )

    return f"""# Научно-технический отчет по моделированию инфракрасного датчика и генерации синтетических изображений

## 1. Введение

Моделирование инфракрасного датчика в данном репозитории выполнялось для оценки того, как температурное состояние наблюдаемой сцены преобразуется в цифровое изображение, пригодное для анализа температурных аномалий. Такая постановка важна для систем мониторинга на базе БПЛА, поскольку реальные натурные измерения трудно проводить во всем диапазоне температур, геометрий наблюдения, уровней шума, дефектности матрицы и режимов коррекции. Синтетические инфракрасные данные в этой работе рассматриваются как управляемый источник испытательных кадров, у которых известны исходные параметры сцены и аппаратной модели.

Практическая область применения модели связана с обнаружением локальных температурных аномалий. В экспериментах ниже рассматриваются как однородные сцены, необходимые для калибровки и анализа физической цепочки, так и сцены с локальной прямоугольной аномалией на фоне 300 K. Отчет построен только на фактически сохраненных данных в папках `results/1.x` и `results/2.x`; проблемные места, включая отрицательные ADC-коды в ранних тестах и необходимость калиброванного ADC-окна в серии 2.x, явно зафиксированы.

## 2. Описание математической модели

Физическая часть модели начинается с расчета излучения черного тела в инфракрасном диапазоне. Спектральная плотность излучения определяется законом Планка:

$$
L_\\lambda(T)=\\frac{{2hc^2}}{{\\lambda^5}}\\frac{{1}}{{\\exp\\left(\\frac{{hc}}{{\\lambda kT}}\\right)-1}}.
$$

В программной реализации `Blackbody` интегрирует излучение в диапазоне длин волн 8--14 мкм и рассчитывает мощность, приходящуюся на чувствительную площадь пикселя. В упрощенной форме мощность центрального пикселя можно представить как

$$
P=A\\Omega\\cos(\\varphi_s)\\cos(\\varphi_r)\\int_{{\\lambda_1}}^{{\\lambda_2}}L_\\lambda(T)\\,d\\lambda,
$$

где $A$ -- чувствительная площадь пикселя, $\\Omega$ -- проектированный телесный угол, $\\varphi_s$ и $\\varphi_r$ -- углы источника и приемника. Оптическая модель формирует пространственное распределение мощности по матрице. Для пикселя с координатами $(x,y)$ относительно центра используется множитель вида

$$
P(x,y)=P_0\\left(\\frac{{f}}{{\\sqrt{{f^2+x^2+y^2}}}}\\right)^4,
$$

что соответствует естественному падению освещенности по полю кадра. Эта стадия важна, поскольку даже идеально однородная температурная сцена после оптики становится неоднородной на матрице.

Далее модель `Bolometers` расширяет активную область до полной структуры с blind и boundary пикселями, добавляет вклад собственного излучения корпуса камеры и задает разброс параметров микроболометров: сопротивления, теплопроводности и теплоемкости. Стадия `Readout` преобразует тепловое воздействие в аналоговый сигнал, решая нелинейное уравнение отклика болометра. Для серии 2.x применен векторизованный расчет, повторяющий структуру уравнений `Readout`, но исключающий чрезвычайно медленное решение `fsolve` для каждого пикселя. После этого сигнал квантуется в `ADC`. В ранних тестах использовался исходный вариант, который выявил отрицательные цифровые коды; в серии 2.x использовано калиброванное аналоговое окно перед квантованием, что исключило отрицательные коды и полное насыщение. Финальная стадия `NUC` выполняет двухточечную линейную коррекцию:

$$
Y_{{corr}}=aY+b.
$$

Коэффициенты $a$ и $b$ вычисляются по двум однородным калибровочным кадрам. В экспериментах оценивалось влияние дробной разрядности этих коэффициентов.

## 3. Архитектура моделирования

Архитектура репозитория соответствует последовательному pipeline:

```text
Temperature scene
      |
      v
Blackbody -> Optics -> Bolometers -> Readout -> ADC -> NUC
      |          |           |          |       |      |
   P(T)      P(x,y)     FPN/noise   analog   digital corrected
```

`Blackbody` задает энергетический вход модели, `Optics` распределяет мощность по фокальной плоскости, `Bolometers` вводит физические параметры матрицы и неравномерность, `Readout` моделирует аналоговую цепь, `ADC` формирует цифровой код, а `NUC` уменьшает фиксированную пространственную неоднородность. Ранние запуски 1.x использовались как диагностические проверки отдельных стадий. Полноценные серии 2.x содержат уже подготовленные экспериментальные данные, графики и таблицы для анализа.

## 4. Анализ запусков 1.x

### 4.1. Запуск 1.0

В запуске 1.0 проверялась первая физическая стадия модели: зависимость мощности, приходящейся на центральный пиксель, от температуры черного тела. Температура изменялась от 280 K до 420 K с шагом 20 K. Минимальная рассчитанная мощность составила {fmt(r10['power_min_W'])} W, максимальная -- {fmt(r10['power_max_W'])} W; отношение максимума к минимуму равно {fmt(r10['power_ratio_max_to_min'])}. Это подтверждает монотонный рост in-band излучения в диапазоне 8--14 мкм.

![Мощность на центральном пикселе](images/run_1_0_power_vs_temperature.png)

![Гистограмма мощностей](images/run_1_0_power_histogram.png)

### 4.2. Запуск 1.1

Запуск 1.1 включал стадии `Blackbody -> Optics` при разрешении 64x48 и температурах 300, 350 и 400 K. Для 400 K центральная мощность составила {fmt(r11[-1]['center_power_W'])} W, угловая -- {fmt(r11[-1]['corner_power_W'])} W, а отношение угла к центру равно {fmt(r11[-1]['corner_to_center_ratio'])}. Следовательно, уже на оптической стадии равномерная сцена получает пространственный градиент.

![Распределение мощности после оптики](images/run_1_1_heatmap_400K.png)

![Гистограмма мощности после оптики](images/run_1_1_hist_400K.png)

### 4.3. Запуск 1.2

Запуск 1.2 был первой попыткой малого полного конвейера `Blackbody -> Optics -> Bolometers -> Readout -> ADC` на матрице 8x6. Он завершился ошибкой на стадии `Readout.process`. Причина состояла в том, что `scipy.optimize.fsolve()` возвращал одноэлементный массив, а исходный код присваивал его скалярной ячейке NumPy. Этот результат важен инженерно: он показал, что исходный readout-код зависит от старого поведения NumPy или требует явного извлечения скаляра.

![Диагностическая ошибка запуска 1.2](images/run_1_2_failure.png)

### 4.4. Запуск 1.3

После перехода к совместимому окружению NumPy 1.26.4 малый полный конвейер до ADC выполнился. При 300 K среднее значение ADC составило {fmt(r13[0]['mean_adc'])}, при 320 K -- {fmt(r13[1]['mean_adc'])}. Оба значения отрицательны, что физически некорректно для цифрового кода АЦП. Этот результат не был отброшен: он зафиксировал необходимость нижнего насыщения или смещения перед квантованием.

![ADC-кадр малого полного конвейера](images/run_1_3_adc_320K.png)

![Гистограмма ADC-кодов малого конвейера](images/run_1_3_adc_320K_hist.png)

![Средний ADC-код в запуске 1.3](images/run_1_3_mean_adc.png)

### 4.5. Запуск 1.4

В запуске 1.4 модель `NUC` применялась к кадрам из запуска 1.3. Нефинитных коэффициентов не возникло: для `coef_a` и `coef_b` количество non-finite значений равно {r14[0]['nonfinite_count']} и {r14[1]['nonfinite_count']} соответственно. Однако поскольку входные ADC-кадры имели отрицательный диапазон, скорректированные значения также оставались в нештатном диапазоне. Запуск полезен как проверка математической работоспособности NUC и как подтверждение необходимости корректной модели ADC.

![NUC на малом кадре](images/run_1_4_nuc_320K.png)

![Гистограмма после NUC](images/run_1_4_nuc_320K_hist.png)

## 5. Анализ запусков 2.x

### 5.1. Запуск 2.0

Запуск 2.0 является базовым прогоном однородной сцены через цепочку `Blackbody -> Optics -> Bolometers -> VectorizedReadout -> ADC`. Использовалось разрешение 160x120 и температуры от 280 K до 400 K. В отличие от диагностического запуска 1.3, здесь применено калиброванное ADC-окно, поэтому отрицательные коды и полное насыщение исключены. Средний ADC-код вырос с {fmt(r20[0]['mean'])} при 280 K до {fmt(r20[-1]['mean'])} при 400 K. Доля насыщенных пикселей при 400 K равна {fmt(r20[-1]['saturated_fraction'])}, доля нулевых пикселей при 280 K равна {fmt(r20[0]['zero_fraction'])}.

{table_20}

![ADC-кадр при 400 K](images/run_2_0_adc_400K.png)

![Гистограмма ADC при 400 K](images/run_2_0_adc_400K_hist.png)

![Средний ADC-код от температуры](images/run_2_0_mean_adc_vs_temperature.png)

### 5.2. Запуск 2.1

Запуск 2.1 изолировал влияние оптики. Моделировались три варианта FOV: narrow, medium и wide. При 300 K максимальное падение мощности в углу составило {fmt(worst21['corner_drop_percent'])}% для режима `{worst21['fov_case']}`. Это означает, что пространственная неоднородность возникает до болометрической матрицы и должна рассматриваться как нормальный оптический эффект, а не только как аппаратный дефект.

{table_21}

![Распределение мощности wide FOV](images/run_2_1_p_300K_wide.png)

![Гистограмма мощности wide FOV](images/run_2_1_p_300K_wide_hist.png)

![Падение мощности в углу от FOV](images/run_2_1_corner_drop_by_fov.png)

### 5.3. Запуск 2.2

Запуск 2.2 оценивал fixed pattern noise от разброса параметров микроболометров. Для изоляции FPN оптический градиент был отключен, использовался flat-field режим. Tolerance изменялся от 1e-6 до 1e-4, для каждого уровня выполнялись реализации с тремя seed. Максимальное среднее стандартное отклонение составило {fmt(worst22['std_mean'])} ADC code при tolerance {worst22['tolerance']}.

{table_22}

![ADC-карта при большом tolerance](images/run_2_2_adc_high_tolerance.png)

![Гистограмма ADC при большом tolerance](images/run_2_2_adc_high_tolerance_hist.png)

![FPN STD от tolerance](images/run_2_2_fpn_std_vs_tolerance.png)

### 5.4. Запуск 2.3

Запуск 2.3 исследовал влияние дефектных пикселей. В базовый кадр с фоном 300 K и аномалией 360 K вводились hot/cold дефекты с долей от 0% до 2%. При 0% дефектов средний STD составил {fmt(r23agg[0]['std_mean'])}, а при 2% -- {fmt(r23agg[-1]['std_mean'])}. Среднее число выбросов за пределами mean ± 3 std при 2% дефектов равно {fmt(r23agg[-1]['outlier_count_mean_pm_3std_mean'])}. Следовательно, дефекты влияют не только визуально, но и статистически.

{table_23}

![ADC-карта с 2% дефектов](images/run_2_3_adc_defects_2pct.png)

![Гистограмма ADC с дефектами](images/run_2_3_adc_defects_2pct_hist.png)

![STD от доли дефектов](images/run_2_3_std_vs_defects.png)

### 5.5. Запуск 2.4

Запуск 2.4 оценивал двухточечную коррекцию NUC по калибровочным кадрам 300 K и 360 K. Тестовые температуры составляли 320, 340 и 380 K; коэффициенты проверялись при точности 3, 4, 5, 6 бит и full precision. Для температуры 340 K стандартное отклонение до коррекции было {fmt([row for row in r24 if row['temperature_K'] == 340][0]['std_before'])}. После full precision коррекции оно снизилось до {fmt(best24['std_after'])} ADC code в лучшем случае по всей серии. При 6-битных коэффициентах для 340 K остаточный STD равен {fmt([row for row in r24 if row['temperature_K'] == 340 and row['fractional_bits'] == '6bit'][0]['std_after'])}, а при 3-битных -- {fmt([row for row in r24 if row['temperature_K'] == 340 and row['fractional_bits'] == '3bit'][0]['std_after'])}.

{table_24}

![Сравнение кадра до и после NUC](images/run_2_4_nuc_340K.png)

![Гистограмма после NUC](images/run_2_4_corrected_340K_hist.png)

![STD после NUC от точности коэффициентов](images/run_2_4_std_after_by_precision.png)

### 5.6. Запуск 2.5

Запуск 2.5 является интеграционным экспериментом с температурной аномалией. Фон имел температуру 300 K, а аномалия принимала значения 320, 340, 360 и 400 K. Проверялись размеры small, medium и large, два уровня tolerance и два режима: до и после NUC. Метрика SNR-like вычислялась как

$$
SNR=\\frac{{\\mu_{{anom}}-\\mu_{{bg}}}}{{\\sigma_{{bg}}}},
$$

при этом для устойчивости после почти идеального выравнивания фона использовался шумовой пол 1 ADC code. Максимальная SNR-like оценка составила {fmt(best25['snr_like'])} для аномалии {best25['anomaly_temperature_K']} K, размера `{best25['anomaly_size']}`, tolerance `{best25['tolerance_case']}` и состояния `{best25['state']}`.

{table_25}

![Карта температурной сцены](images/run_2_5_scene_large_400K.png)

![Аномалия до и после NUC](images/run_2_5_anomaly_400K.png)

![Гистограмма corrected-кадра](images/run_2_5_corrected_400K_hist.png)

![Контраст аномалии от температуры](images/run_2_5_contrast_vs_temperature.png)

## 6. Сравнительный анализ

Сравнение всех запусков показывает, что температура является главным фактором, определяющим уровень полезного сигнала. В базовом прогоне 2.0 средний ADC-код изменился с {fmt(r20[0]['mean'])} до {fmt(r20[-1]['mean'])}, то есть почти на весь рабочий диапазон калиброванного АЦП. Этот результат согласуется с ростом мощности черного тела в запуске 1.0.

Оптика влияет не на общий уровень температуры как таковой, а на пространственную форму кадра. В запуске 2.1 wide FOV дал падение угловой мощности до {fmt(worst21['corner_drop_percent'])}%, что существенно для задач обнаружения объектов на периферии кадра. В отличие от оптики, разброс параметров болометров формирует случайно-фиксированную структуру. В запуске 2.2 рост tolerance до 1e-4 дал средний STD {fmt(worst22['std_mean'])} ADC code, что сопоставимо с заметными изменениями яркости внутри кадра.

Дефектные пиксели проявились как более локальная, но статистически значимая деградация: при росте доли дефектов до 2% стандартное отклонение увеличилось с {fmt(r23agg[0]['std_mean'])} до {fmt(r23agg[-1]['std_mean'])}, а число выбросов достигло {fmt(r23agg[-1]['outlier_count_mean_pm_3std_mean'])}. NUC оказалась наиболее сильным корректирующим механизмом. В запуске 2.4 для full precision остаточная неоднородность снизилась до долей ADC-кода, а зависимость от разрядности коэффициентов показала, что грубая 3-битная аппроксимация оставляет заметный остаточный FPN.

В прикладном запуске 2.5 NUC повышала наблюдаемость аномалий за счет уменьшения фоновой неоднородности. Однако при интерпретации SNR-like после NUC необходимо учитывать шумовой пол, иначе идеально выровненный фон приводит к математически бесконечным или завышенным значениям отношения сигнал/шум. Это не является ошибкой модели, но является ограничением самой метрики при отсутствии отдельной модели случайного временного шума.

## 7. Выводы

Проведенные эксперименты подтверждают работоспособность цепочки моделирования инфракрасного датчика и показывают физически осмысленную связь между температурой сцены, мощностью излучения, оптическим распределением, болометрической неоднородностью, readout-сигналом и цифровым кадром. Диагностические запуски 1.x выявили важные инженерные проблемы: несовместимость исходного readout-кода с современным поведением `fsolve`/NumPy и возможность отрицательных ADC-кодов без корректного смещения или насыщения. Эти наблюдения были учтены при построении серии 2.x.

Полноценные запуски 2.x показали, что температура и оптика задают базовый радиометрический и пространственный профиль кадра, разброс болометров формирует fixed pattern noise, дефектные пиксели увеличивают число выбросов, а двухточечная NUC существенно снижает неоднородность. Интеграционный эксперимент с температурной аномалией подтвердил, что синтетические данные пригодны для оценки обнаружимости аномалий при варьировании температуры, размера, tolerance и режима коррекции.

Ограничения модели также должны учитываться в дипломной интерпретации. В серии 2.x использован векторизованный readout как инженерно необходимое ускорение, а ADC-квантование выполнялось через калиброванное окно, чтобы избежать полного насыщения. Кроме того, SNR-like в отсутствие отдельного временного шума требует введения шумового пола. Несмотря на эти ограничения, модель является пригодной основой для генерации синтетических ИК-изображений и для инженерного анализа системы мониторинга температурных аномалий.

## Сводная таблица экспериментов

{summary_table}
"""


def build_docx_from_markdown(markdown_path: Path, docx_path: Path) -> bool:
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception:
        return False

    def split_table_row(row: str) -> list[str]:
        return [cell.strip() for cell in row.strip().strip("|").split("|")]

    document = Document()
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code_block = False
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            document.add_paragraph(stripped)
            index += 1
            continue
        if stripped.startswith("|") and not in_code_block:
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            table_rows = [split_table_row(row) for row in table_lines]
            table_rows = [row for row in table_rows if not all(set(cell) <= {"-"} for cell in row)]
            if table_rows:
                cols = len(table_rows[0])
                table = document.add_table(rows=1, cols=cols)
                table.style = "Table Grid"
                for col_idx, cell in enumerate(table_rows[0]):
                    table.rows[0].cells[col_idx].text = cell
                for table_row in table_rows[1:]:
                    cells = table.add_row().cells
                    for col_idx in range(cols):
                        cells[col_idx].text = table_row[col_idx] if col_idx < len(table_row) else ""
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
        elif stripped.startswith("![") and "](" in stripped and stripped.endswith(")"):
            image_path = stripped.split("](", 1)[1][:-1]
            full_path = DOCS / image_path
            if full_path.exists():
                document.add_picture(str(full_path), width=Inches(5.8))
        elif stripped.startswith("$$"):
            document.add_paragraph(stripped)
        elif stripped:
            document.add_paragraph(stripped)
        index += 1
    document.save(docx_path)
    return True


def main() -> None:
    images = build_images()
    data = collect_data()
    summary_rows = build_experiment_summary(data)
    write_csv(
        DOCS / "experiment_summary.csv",
        summary_rows,
        ["run", "key_parameters", "key_metrics", "interpretation"],
    )
    report = report_text(data, images, summary_rows)
    report_path = DOCS / "report.md"
    report_path.write_text(report, encoding="utf-8")
    if build_docx_from_markdown(report_path, DOCS / "report.docx"):
        print("Generated docs/report.docx")
    else:
        print("Skipped docs/report.docx: python-docx is not installed")
    print("Generated docs/report.md")
    print("Generated docs/experiment_summary.csv")
    print(f"Generated {len(list(IMAGES.glob('*.png')))} images in docs/images")


if __name__ == "__main__":
    main()
